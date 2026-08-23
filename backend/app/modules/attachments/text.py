from __future__ import annotations

import codecs
from pathlib import PurePath
from zipfile import ZipFile

from docx import Document

DOCX_MEDIA_TYPE = (
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
)
MAX_DOCX_ARCHIVE_ENTRIES = 4096
MAX_DOCX_UNCOMPRESSED_BYTES = 16 * 1024 * 1024
MAX_SEARCH_TEXT_BYTES = 1024 * 1024


def _read(source) -> bytes:
    source.seek(0)
    try:
        return source.read(MAX_SEARCH_TEXT_BYTES)
    finally:
        source.seek(0)


def _plain_text(source) -> str:
    data = _read(source).removeprefix(codecs.BOM_UTF8)
    return _utf8_prefix(data)


def _utf8_prefix(data: bytes) -> str:
    decoder = codecs.getincrementaldecoder("utf-8")()
    return decoder.decode(data[:MAX_SEARCH_TEXT_BYTES], final=False)


def _limit_text(text: str) -> str:
    return _utf8_prefix(text.encode("utf-8"))


def _require_safe_docx(source) -> None:
    source.seek(0)
    try:
        with ZipFile(source) as archive:
            entries = archive.infolist()
            if len(entries) > MAX_DOCX_ARCHIVE_ENTRIES:
                raise ValueError("DOCX contains too many archive entries")
            if sum(entry.file_size for entry in entries) > MAX_DOCX_UNCOMPRESSED_BYTES:
                raise ValueError("DOCX exceeds the uncompressed size limit")
    finally:
        source.seek(0)


def _docx_text(source) -> str:
    _require_safe_docx(source)
    source.seek(0)
    try:
        document = Document(source)
    finally:
        source.seek(0)
    paragraphs = [paragraph.text for paragraph in document.paragraphs]
    cells = [
        cell.text
        for table in document.tables
        for row in table.rows
        for cell in row.cells
    ]
    text = "\n".join(text for text in (*paragraphs, *cells) if text)
    return _limit_text(text)


def extract_search_text(upload) -> str:
    media_type = upload.content_type or "application/octet-stream"
    suffix = PurePath(upload.filename or "").suffix.lower()
    if media_type.startswith("text/"):
        return _plain_text(upload.file)
    if media_type == DOCX_MEDIA_TYPE or suffix == ".docx":
        return _docx_text(upload.file)
    return ""
