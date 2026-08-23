from __future__ import annotations

from io import BytesIO
from pathlib import Path

from docx import Document
from docx.document import Document as DocxDocument
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Mm, Pt
from docx.text.paragraph import Paragraph
from docx.text.run import Run

from app.modules.cases.document_schema import validate_prosemirror_document
from app.modules.documents.fonts import embed_title_font
from app.modules.documents.styles import (
    BODY_FONT,
    BODY_SIZE,
    HEADING_FONT,
    HEADING_SIZE,
    TITLE_FONT,
    TITLE_SIZE,
    configure_document,
    format_body,
    format_heading,
    format_title,
    set_run_font,
)

LOGO_PATH = (
    Path(__file__).with_name("assets") / "shanghai-university-horizontal-logo.png"
)


def _mark_types(node: dict) -> set[str]:
    return {mark.get("type", "") for mark in node.get("marks", [])}


def _format_marks(run: Run, node: dict) -> None:
    marks = _mark_types(node)
    run.bold = "bold" in marks
    run.italic = "italic" in marks
    run.font.strike = "strike" in marks


def _add_inlines(paragraph: Paragraph, nodes: list[dict], font: str, size: int) -> None:
    for node in nodes:
        if node.get("type") == "hardBreak":
            paragraph.add_run().add_break()
            continue
        if node.get("type") == "text":
            run = paragraph.add_run(str(node.get("text", "")))
            set_run_font(run, font, size)
            _format_marks(run, node)
        elif node.get("content"):
            _add_inlines(paragraph, node["content"], font, size)


def _add_title(document: DocxDocument, title: str) -> None:
    paragraph = document.add_paragraph(style="Title")
    format_title(paragraph)
    run = paragraph.add_run(title)
    set_run_font(run, TITLE_FONT, TITLE_SIZE)


def _add_logo(document: DocxDocument) -> None:
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_after = Pt(10)
    shape = paragraph.add_run().add_picture(str(LOGO_PATH), width=Mm(60))
    shape._inline.docPr.set("descr", "上海大学校徽")


def _add_heading(document: DocxDocument, node: dict) -> None:
    level = min(3, max(1, int(node.get("attrs", {}).get("level", 1))))
    paragraph = document.add_paragraph(style=f"Heading {level}")
    format_heading(paragraph)
    _add_inlines(paragraph, node.get("content", []), HEADING_FONT, HEADING_SIZE)


def _add_paragraph(
    document: DocxDocument, node: dict, style: str | None = None
) -> Paragraph:
    paragraph = document.add_paragraph(style=style)
    format_body(paragraph, indented=style is None)
    _add_inlines(paragraph, node.get("content", []), BODY_FONT, BODY_SIZE)
    return paragraph


def _bind_numbering(paragraph: Paragraph, num_id: int) -> None:
    properties = paragraph._p.get_or_add_pPr().get_or_add_numPr()
    properties.get_or_add_ilvl().val = 0
    properties.get_or_add_numId().val = num_id


def _new_numbering(document: DocxDocument, style: str, start: int) -> int:
    style_num = document.styles[style].element.pPr.numPr.numId.val
    numbering = document.part.numbering_part.element
    abstract_num = numbering.num_having_numId(style_num).abstractNumId.val
    instance = numbering.add_num(abstract_num)
    if start != 1:
        instance.add_lvlOverride(ilvl=0).add_startOverride(val=start)
    return instance.numId


def _add_list(
    document: DocxDocument, node: dict, style: str, num_id: int | None = None
) -> None:
    for item in node.get("content", []):
        children = item.get("content", [])
        for index, child in enumerate(children):
            if child.get("type") == "paragraph":
                paragraph = _add_paragraph(
                    document, child, style if index == 0 else None
                )
                if index == 0 and num_id is not None:
                    _bind_numbering(paragraph, num_id)
            else:
                _add_node(document, child)


def _add_blockquote(document: DocxDocument, node: dict) -> None:
    for child in node.get("content", []):
        if child.get("type") == "paragraph":
            _add_paragraph(document, child, "Quote")
        else:
            _add_node(document, child)


def _add_node(document: DocxDocument, node: dict) -> None:
    kind = node.get("type")
    if kind == "heading":
        _add_heading(document, node)
    elif kind == "paragraph":
        _add_paragraph(document, node)
    elif kind == "blockquote":
        _add_blockquote(document, node)
    elif kind == "bulletList":
        _add_list(document, node, "List Bullet")
    elif kind == "orderedList":
        start = int(node.get("attrs", {}).get("start", 1))
        _add_list(
            document,
            node,
            "List Number",
            _new_numbering(document, "List Number", start),
        )


def build_case_docx(case: dict) -> bytes:
    validate_prosemirror_document(case["document"])
    document = Document()
    configure_document(document)
    document.core_properties.title = str(case["title"])
    _add_logo(document)
    _add_title(document, str(case["title"]))
    for node in case["document"].get("content", []):
        _add_node(document, node)
    buffer = BytesIO()
    document.save(buffer)
    return embed_title_font(buffer.getvalue())
