from __future__ import annotations

from io import BytesIO
from zipfile import ZIP_DEFLATED, ZIP_STORED, ZipFile

import pytest
from docx import Document
from fastapi.testclient import TestClient

from app.modules.attachments.text import DOCX_MEDIA_TYPE, MAX_SEARCH_TEXT_BYTES


def _login(client: TestClient, username: str, password: str) -> dict:
    response = client.post(
        "/api/auth/login",
        json={"username": username, "password": password},
    )
    assert response.status_code == 200
    return response.json()


def _revision(client: TestClient) -> int:
    return client.get("/api/cases/c-draft-1").json()["revision"]


def _fail_remove(_blob_id: str) -> None:
    raise OSError("object store unavailable")


def _upload_attachment(client: TestClient, headers: dict, name: str, body: bytes):
    return client.post(
        "/api/cases/c-draft-1/attachments",
        headers=headers,
        data={"accessLevel": "private", "revision": _revision(client)},
        files={"file": (name, body, "text/plain")},
    ).json()


def _delete_attachment(client: TestClient, headers: dict, attachment_id: str):
    return client.delete(
        f"/api/cases/c-draft-1/attachments/{attachment_id}",
        headers=headers,
        params={"revision": _revision(client)},
    )


def _large_docx() -> bytes:
    source = BytesIO()
    document = Document()
    document.add_paragraph("附件大文档雪松词")
    document.save(source)
    output = BytesIO(source.getvalue())
    with ZipFile(output, "a", ZIP_STORED) as archive:
        archive.writestr("word/media/large.bin", b"x" * (1024 * 1024 + 1))
    return output.getvalue()


def _docx_with_oversized_archive() -> bytes:
    source = BytesIO()
    document = Document()
    document.add_paragraph("压缩文档")
    document.save(source)
    with ZipFile(source, "a", ZIP_DEFLATED) as archive:
        archive.writestr("word/media/bomb.bin", b"x" * (17 * 1024 * 1024))
    return source.getvalue()


def test_author_can_upload_and_list_a_public_attachment(client: TestClient) -> None:
    auth = _login(client, "user", "user123")
    response = client.post(
        "/api/cases/c-draft-1/attachments",
        headers={"X-CSRF-Token": auth["csrfToken"]},
        data={"accessLevel": "public", "revision": _revision(client)},
        files={"file": ("课堂证据.txt", b"source evidence", "text/plain")},
    )

    assert response.status_code == 201
    attachment = response.json()
    assert attachment["name"] == "课堂证据.txt"
    assert attachment["accessLevel"] == "public"
    assert attachment["mediaType"] == "text/plain"
    assert attachment["size"] == 15
    assert "blobId" not in attachment
    assert "objectKey" not in attachment
    assert client.get("/api/cases/c-draft-1/attachments").json() == [attachment]


def test_submission_freezes_extracted_attachment_text_without_exposing_it(
    client: TestClient,
) -> None:
    auth = _login(client, "user", "user123")
    headers = {"X-CSRF-Token": auth["csrfToken"]}
    attachment = _upload_attachment(
        client,
        headers,
        "检索附件.txt",
        "附件唯一雪松词".encode(),
    )
    assert "searchText" not in attachment
    database = client.app.state.database
    stored = database.attachments.find_one({"id": attachment["id"]})
    assert stored["searchText"] == "附件唯一雪松词"
    submitted = client.post(
        "/api/cases/c-draft-1/lifecycle",
        headers=headers,
        json={"command": "submit", "revision": _revision(client)},
    ).json()
    version = database.case_versions.find_one({"id": submitted["version"]["id"]})
    assert version["attachments"][0]["searchText"] == "附件唯一雪松词"


def test_author_can_upload_a_valid_docx_larger_than_search_text_limit(
    client: TestClient,
) -> None:
    auth = _login(client, "user", "user123")
    response = client.post(
        "/api/cases/c-draft-1/attachments",
        headers={"X-CSRF-Token": auth["csrfToken"]},
        data={"accessLevel": "private", "revision": _revision(client)},
        files={"file": ("large.docx", _large_docx(), DOCX_MEDIA_TYPE)},
    )

    assert response.status_code == 201
    stored = client.app.state.database.attachments.find_one(
        {"id": response.json()["id"]}
    )
    assert stored["searchText"] == "附件大文档雪松词"


def test_upload_rejects_a_malformed_docx(client: TestClient) -> None:
    auth = _login(client, "user", "user123")
    response = client.post(
        "/api/cases/c-draft-1/attachments",
        headers={"X-CSRF-Token": auth["csrfToken"]},
        data={"accessLevel": "private", "revision": _revision(client)},
        files={"file": ("broken.docx", b"not a docx", DOCX_MEDIA_TYPE)},
    )

    assert response.status_code == 422
    assert response.json() == {"detail": "附件文本无法解析"}


def test_upload_rejects_a_docx_above_the_uncompressed_budget(
    client: TestClient,
) -> None:
    auth = _login(client, "user", "user123")
    body = _docx_with_oversized_archive()
    assert len(body) < 128 * 1024
    response = client.post(
        "/api/cases/c-draft-1/attachments",
        headers={"X-CSRF-Token": auth["csrfToken"]},
        data={"accessLevel": "private", "revision": _revision(client)},
        files={"file": ("bomb.docx", body, DOCX_MEDIA_TYPE)},
    )

    assert response.status_code == 422
    assert response.json() == {"detail": "附件文本无法解析"}


def test_plain_text_stops_before_a_split_utf8_character(client: TestClient) -> None:
    auth = _login(client, "user", "user123")
    body = b"a" * (MAX_SEARCH_TEXT_BYTES - 1) + "雪".encode()
    response = client.post(
        "/api/cases/c-draft-1/attachments",
        headers={"X-CSRF-Token": auth["csrfToken"]},
        data={"accessLevel": "private", "revision": _revision(client)},
        files={"file": ("large.txt", body, "text/plain")},
    )

    assert response.status_code == 201
    stored = client.app.state.database.attachments.find_one(
        {"id": response.json()["id"]}
    )
    assert stored["searchText"] == "a" * (MAX_SEARCH_TEXT_BYTES - 1)


def test_author_can_download_an_attachment(client: TestClient) -> None:
    auth = _login(client, "user", "user123")
    created = client.post(
        "/api/cases/c-draft-1/attachments",
        headers={"X-CSRF-Token": auth["csrfToken"]},
        data={"accessLevel": "private", "revision": _revision(client)},
        files={"file": ("evidence.txt", b"download me", "text/plain")},
    ).json()

    response = client.get(f"/api/cases/c-draft-1/attachments/{created['id']}/content")

    assert response.status_code == 200
    assert response.content == b"download me"
    assert response.headers["content-type"] == "text/plain; charset=utf-8"
    assert "evidence.txt" in response.headers["content-disposition"]


def test_author_can_delete_a_draft_attachment(client: TestClient) -> None:
    auth = _login(client, "user", "user123")
    headers = {"X-CSRF-Token": auth["csrfToken"]}
    created = _upload_attachment(client, headers, "remove.txt", b"remove me")
    response = _delete_attachment(client, headers, created["id"])

    assert response.status_code == 204
    assert client.get("/api/cases/c-draft-1/attachments").json() == []
    assert (
        client.get(
            f"/api/cases/c-draft-1/attachments/{created['id']}/content"
        ).status_code
        == 404
    )


def test_delete_commits_when_orphan_blob_cleanup_is_temporarily_unavailable(
    client: TestClient, monkeypatch: pytest.MonkeyPatch
) -> None:
    auth = _login(client, "user", "user123")
    headers = {"X-CSRF-Token": auth["csrfToken"]}
    created = _upload_attachment(client, headers, "orphan.txt", b"orphan")
    store = client.app.state.blob_store
    monkeypatch.setattr(store, "remove", _fail_remove)
    response = _delete_attachment(client, headers, created["id"])

    assert response.status_code == 204
    assert client.get("/api/cases/c-draft-1/attachments").json() == []


def test_admin_can_read_but_cannot_mutate_an_author_draft(client: TestClient) -> None:
    author = _login(client, "user", "user123")
    created = client.post(
        "/api/cases/c-draft-1/attachments",
        headers={"X-CSRF-Token": author["csrfToken"]},
        data={"accessLevel": "private", "revision": _revision(client)},
        files={"file": ("review.txt", b"review", "text/plain")},
    ).json()
    admin = _login(client, "admin", "admin123")
    headers = {"X-CSRF-Token": admin["csrfToken"]}

    assert client.get("/api/cases/c-draft-1/attachments").status_code == 200
    assert (
        client.get(
            f"/api/cases/c-draft-1/attachments/{created['id']}/content"
        ).status_code
        == 200
    )
    assert (
        client.delete(
            f"/api/cases/c-draft-1/attachments/{created['id']}",
            headers=headers,
            params={"revision": _revision(client)},
        ).status_code
        == 403
    )


def test_upload_rejects_a_file_above_128_mib(
    client: TestClient, monkeypatch: pytest.MonkeyPatch
) -> None:
    monkeypatch.setattr("app.modules.attachments.service.MAX_ATTACHMENT_BYTES", 8)
    auth = _login(client, "user", "user123")
    response = client.post(
        "/api/cases/c-draft-1/attachments",
        headers={"X-CSRF-Token": auth["csrfToken"]},
        data={"accessLevel": "campus", "revision": _revision(client)},
        files={"file": ("large.bin", b"123456789", "application/octet-stream")},
    )

    assert response.status_code == 413
    assert response.json() == {"detail": "附件不能超过 128MiB"}
