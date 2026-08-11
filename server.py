#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
思政教学案例智能平台 — 服务端

职责：
1. 托管 app/ 下的前端；
2. /api/ai/chat      —— OpenAI 兼容 Chat 接口服务端代理（API Key 只存在服务端，支持 SSE 流式）；
3. /api/web-search   —— AnySearch/Tavily 联网检索（密钥只存在服务端）；
4. /api/fetch-url    —— 公开 URL 采集（生成可引用的内容副本，原始 URL 保留）；
5. /api/export-docx  —— 将案例及成套教学材料导出为 .docx（页脚注入追踪元数据）；
6. /api/auth/login   —— 演示登录：账号 ID 换 HMAC token（12h，无密码，ADR 0009）；
7. /api/files        —— 资料文件库：列表/上传/下载/删除，文件落盘 files/，按密级服务端强制；
   上传 md/txt/docx 自动抽取纯文本，GET /api/files/{fid}/text 在线查看；
8. /api/knowledge    —— 知识库在线导入（admin，markdown 按 # 章/### 节解析）与公开查询；
9. /api/constants    —— 健康检查与能力开关；
10. /api/search      —— 服务端语料检索（BM25 + 中文 bigram，惰性索引，按用户密级过滤素材，
   案例按身份过滤；命中带 materialId/sec 结构路径供切片深链，ADR 0010；
   素材语料只索引 status=正常 的，候选不进生成语料）；
11. /api/ai/agent    —— 统一 AI 入口（MoA 多智能体编排：主Agent/资料管理员/写作手/内容审校员，SSE）；
12. /api/cases 等    —— 案例业务闭环（SQLite 持久层 db.py）：案例 CRUD、提交/撤回/审核流转留痕、
   批注与回复线程、版本快照（含与上一版 diffSummary）/回滚、收藏、点赞；
   教师提交前主动调用 AI 自检；退回/要求补充必须带 reasonType；
   /api/admin/review-ledger 输出被退回表达台账；首启自动灌入 files/cases_seed.json；
13. /api/materials 等 —— 素材登记闭环（SQLite materials 表，ADR 0003/0011）：列表/详情/采集入库闸
   （URL 查重 + 相似度查重 + 必填校验，新素材一律先落候选）、admin 治理 PATCH 与批量 PATCH、
   素材收藏、recommendFor 上下文推荐、recentCitedBy 最近引用、来源健康检查；
   首启自动灌入 files/materials_seed.json，被引计数由案例写入时统一重算；
14. /api/admin/watch/* —— 自动盯源（WP5）：栏目源 CRUD 与启停、每小时定时扫描 + 手动触发
   （抓栏目页 → AI 抽取候选条目并做关键词过滤，AI 不可用跳过该源并记录 → URL+标题指纹双重去重
   → 落 watch_items 待审），候选卡入库（走同一入库闸落素材候选，同事件报道作多方验证附注）/忽略；
   版权默认策略：只存标题+摘要+原文链接+元数据，不抓全文；
15. /api/contributions、/api/my/impact —— 众筹雏形（WP5）：素材链接（link）与知识点-素材关联
   （kn_link）贡献先审后发（kn_link 通过后体现在 recommendFor 打分里），仅本人可见自己的贡献；
   /api/my/impact 聚合素材贡献被引次数与案例被收藏/被点赞数；
16. /api/my/prefs（GET/PUT）—— 教师显式生成偏好（WP4b，user_prefs 表：篇幅/风格/禁用词/常用主题，
   只来自教师亲手填写，四项全空即清空）；写作手生成时注入 prompt，禁用词命中在结果中标 risk 警示。
17. /api/onlyoffice/* —— OnlyOffice 主编辑器（WP6）：DS docker 部署（docker-compose.yml，8081，JWT）。
   docx 为编辑真源（files/cases/{caseId}.docx，无则启动时从 blocks 迁移生成）；config 端点发 JWT 签名
   的 DocEditor 配置（key=caseId-docxVer，内容变化即变）；callback 接收 DS 保存（status 2/6 下载落盘
   并回写 blocks、bump docxVer）；insert 端点把 Copilot 采纳内容以 track-changes（w:ins/w:del）写入；
   版本快照复制 docx 到 files/cases/versions/，回滚重建 docx 强制 DS 重载。
18. /api/graph/*    —— Neo4j 知识图谱（WP7，docker-compose neo4j 服务，graph_store.py）：
   节点 Chapter/Knowledge/Case/Material/Tag，边 CITES/MENTIONS/BELONGS/HAS_TAG；
   启动时从 SQLite 全量灌库（LLM 实体增强后台补 MENTIONS 边，AI 不可用则跳过），
   db.py 写路径钩子增量同步，/api/admin/reseed 后自动重建；
   overview（全库轻量图）/ego（两跳子图）/reverse（知识点/标签反查）/qa（全局问答，
   AI 不可用时降级返回子图统计）+ POST /api/admin/graph/rebuild；
   Neo4j 不可达时以上接口返回 degraded 提示，其他功能不受影响。

仅依赖标准库 + python-docx（图谱需 .venv 的 neo4j 驱动）。
运行：.venv/bin/python server.py [port]（/usr/bin/python3 运行时图谱功能关闭）
"""
import base64
import difflib
import hashlib
import hmac
import json
import math
import os
import queue
import re
import socket
import ssl
import sys
import threading
import time
import urllib.error
import urllib.parse
import urllib.request
import uuid
from collections import Counter
from concurrent.futures import ThreadPoolExecutor
from http.server import BaseHTTPRequestHandler, ThreadingHTTPServer

from db import CaseDB
import db
from graph_store import Graph as GraphStore

ROOT = os.path.dirname(os.path.abspath(__file__))
STATIC_DIR = os.path.join(ROOT, "app")


def load_env(path):
    cfg = {}
    if not os.path.exists(path):
        return cfg
    with open(path, "r", encoding="utf-8") as f:
        for line in f:
            line = line.strip()
            if not line or line.startswith("#") or "=" not in line:
                continue
            k, v = line.split("=", 1)
            cfg[k.strip()] = v.strip().strip('"').strip("'")
    for k in list(cfg):  # 环境变量优先于 .env（临时开关/降级演练用）
        if k in os.environ:
            cfg[k] = os.environ[k]
    return cfg


ENV = load_env(os.path.join(ROOT, ".env"))
AI_BASE_URL = ENV.get("AI_BASE_URL", "").rstrip("/")
AI_API_KEY = ENV.get("AI_API_KEY", "")
AI_MODELS = [m.strip() for m in ENV.get("AI_MODELS", "").split(",") if m.strip()]
AI_DEFAULT_MODEL = ENV.get("AI_DEFAULT_MODEL", "qwen-plus")
AI_TIMEOUT = int(ENV.get("AI_TIMEOUT_SECONDS", "60") or 60)
AI_REVIEW_ENABLED = ENV.get("AI_REVIEW_ENABLED", "").lower() in ("1", "true", "yes")
TAVILY_API_KEY = ENV.get("TAVILY_API_KEY", "")
ANYSEARCH_API_KEY = ENV.get("ANYSEARCH_API_KEY", "")
# MoA 多智能体编排模型（.env 可选覆盖；未配置时用以下默认值。写作手主候选 = AI_DEFAULT_MODEL）
MOA_MODEL_ORCHESTRATOR = ENV.get("MOA_MODEL_ORCHESTRATOR", "") or "qwen-plus"
MOA_MODEL_LIBRARIAN = ENV.get("MOA_MODEL_LIBRARIAN", "") or "qwen-plus"
MOA_MODEL_WRITER_ALT = ENV.get("MOA_MODEL_WRITER_ALT", "") or "qwen-plus"
# 注意：模型网关对大小写敏感，R1 的可用名是小写 "deepseek-r1"（大写会被 400 拒绝）
MOA_MODEL_REVIEWER = ENV.get("MOA_MODEL_REVIEWER", "") or "deepseek-r1"

# Neo4j 图谱库（WP7）；未配置或缺 neo4j 驱动（用 .venv/bin/python 运行）时图谱功能关闭
GRAPH = GraphStore(ENV.get("NEO4J_URI", ""), ENV.get("NEO4J_USER", "neo4j"),
                   ENV.get("NEO4J_PASSWORD", ""))

UA = "CaseLibrary/1.0 (+shanghai-university-sizheng-case-library)"

# ------------------------------------------------------------ 演示鉴权与文件库（ADR 0009）
APP_SECRET = ENV.get("APP_SECRET", "")
FILES_DIR = os.path.join(ROOT, "files")
INDEX_FILE = os.path.join(FILES_DIR, "index.json")
USERS_FILE = os.path.join(FILES_DIR, "users.json")
KNOWLEDGE_FILE = os.path.join(FILES_DIR, "knowledge.json")
CASES_SEED_FILE = os.path.join(FILES_DIR, "cases_seed.json")
MATERIALS_SEED_FILE = os.path.join(FILES_DIR, "materials_seed.json")
CASE_ATTACHMENTS_DIR = os.path.join(FILES_DIR, "case-attachments")
LEXICON_FILE = os.path.join(FILES_DIR, "review_lexicon.json")
COUNTEREXAMPLES_FILE = os.path.join(FILES_DIR, "review_counterexamples.json")
SQLITE_DB_PATH = os.path.join(ROOT, ENV.get("SQLITE_DB_PATH", "./data/cases.db"))
CASEDB = None  # main() 启动时初始化（SQLite 业务库）
TOKEN_TTL_SECONDS = 12 * 3600
MAX_UPLOAD_BYTES = 20 * 1024 * 1024
LEVEL_NAMES = ["公开", "校内", "私密"]
_INDEX_LOCK = threading.Lock()


def _b64url(data):
    return base64.urlsafe_b64encode(data).decode("ascii").rstrip("=")


def _b64url_decode(s):
    return base64.urlsafe_b64decode(s + "=" * (-len(s) % 4))


def make_token(uid):
    payload = _b64url(json.dumps(
        {"uid": uid, "exp": int(time.time()) + TOKEN_TTL_SECONDS},
        separators=(",", ":")).encode("utf-8"))
    sig = _b64url(hmac.new(APP_SECRET.encode("utf-8"), payload.encode("ascii"),
                           hashlib.sha256).digest())
    return payload + "." + sig


def read_token(tok):
    """校验 HMAC token，返回 uid；无效/过期返回 None。"""
    if not APP_SECRET or not tok or "." not in tok:
        return None
    payload, sig = tok.rsplit(".", 1)
    want = _b64url(hmac.new(APP_SECRET.encode("utf-8"), payload.encode("ascii"),
                            hashlib.sha256).digest())
    if not hmac.compare_digest(sig, want):
        return None
    try:
        data = json.loads(_b64url_decode(payload))
    except Exception:
        return None
    if data.get("exp", 0) < time.time():
        return None
    return data.get("uid")


def load_users():
    try:
        with open(USERS_FILE, encoding="utf-8") as f:
            return {u["id"]: u for u in json.load(f)}
    except Exception:
        return {}


def auth_user(handler):
    """从 Authorization: Bearer 解析登录用户；未登录返回 None（按公开级看待）。"""
    ah = handler.headers.get("Authorization") or ""
    uid = read_token(ah[7:].strip()) if ah.startswith("Bearer ") else None
    return load_users().get(uid) if uid else None


def load_index():
    try:
        with open(INDEX_FILE, encoding="utf-8") as f:
            return json.load(f).get("files", [])
    except Exception:
        return []


def save_index(entries):
    tmp = INDEX_FILE + ".tmp"
    with open(tmp, "w", encoding="utf-8") as f:
        json.dump({"files": entries}, f, ensure_ascii=False, indent=1)
    os.replace(tmp, INDEX_FILE)


def api_files_list(user):
    entries = [e for e in load_index() if _file_content_visible(e, user)]
    # 素材登记权威在 SQLite materials 表（/api/materials），这里只回文件索引
    return {
        "ok": True,
        "files": {e["id"]: {"name": e["name"], "size": e.get("size", 0), "textPath": e.get("textPath")} for e in entries},
    }


def api_login(payload):
    if not APP_SECRET:
        return {"ok": False, "error": "服务端未配置 APP_SECRET"}
    users = load_users()
    if not users:
        return {"ok": False, "error": "服务端缺少 files/users.json，请先运行 tools/build_data.py"}
    uid = (payload.get("userId") or "").strip()
    u = users.get(uid)
    if not u:
        return {"ok": False, "error": "账号不存在"}
    return {"ok": True, "token": make_token(uid), "user": u}


def extract_upload_text(full_path, ext):
    """按扩展名抽取纯文本（md/txt 直接读，docx 用 python-docx）；
    不支持的类型或抽取失败返回 None，不影响上传本身。"""
    ext = ext.lower()
    try:
        if ext in (".md", ".markdown", ".txt"):
            with open(full_path, encoding="utf-8") as f:
                return f.read()
        if ext == ".docx":
            from docx import Document
            doc = Document(full_path)
            return "\n".join(p.text for p in doc.paragraphs if p.text.strip())
    except Exception:
        return None
    return None


def api_file_upload(user, payload):
    if not user:
        return {"ok": False, "error": "未登录或登录已过期，请重新切换账号"}, 401
    if not user.get("admin"):
        return {"ok": False, "error": "仅案例管理员可上传资料文件"}, 403
    title = (payload.get("title") or "").strip()
    if not title:
        return {"ok": False, "error": "请填写素材标题"}, 400
    try:
        level = int(payload.get("level"))
    except Exception:
        level = -1
    if level not in (0, 1, 2):
        return {"ok": False, "error": "密级取值无效"}, 400
    name = os.path.basename((payload.get("filename") or "").strip()) or "file.bin"
    try:
        raw = base64.b64decode(payload.get("dataBase64") or "")
    except Exception:
        return {"ok": False, "error": "文件内容编码无效"}, 400
    if not raw:
        return {"ok": False, "error": "文件为空"}, 400
    if len(raw) > MAX_UPLOAD_BYTES:
        return {"ok": False, "error": "文件超过 20MB 上限"}, 400

    fid = "f-up-" + uuid.uuid4().hex[:12]
    ext = re.sub(r"[^A-Za-z0-9.]", "", os.path.splitext(name)[1])[:10]
    rel = "up/" + fid + ext
    os.makedirs(os.path.join(FILES_DIR, "up"), exist_ok=True)
    with open(os.path.join(FILES_DIR, rel), "wb") as f:
        f.write(raw)
    entry = {
        "id": fid, "materialId": "m-up-" + uuid.uuid4().hex[:12],
        "name": name, "path": rel, "size": len(raw), "level": level,
        "seed": False, "title": title,
        "summary": (payload.get("summary") or "").strip(),
        "by": user["id"], "byName": user.get("name", ""),
        "at": time.strftime("%Y-%m-%d %H:%M"),
    }
    # 在线文本抽取：成功则落盘 up/{fid}.txt 并在索引条目上记录 textPath
    text = extract_upload_text(os.path.join(FILES_DIR, rel), ext)
    if text and text.strip():
        text_rel = "up/" + fid + ".txt"
        with open(os.path.join(FILES_DIR, text_rel), "w", encoding="utf-8") as f:
            f.write(text)
        entry["textPath"] = text_rel
    with _INDEX_LOCK:
        entries = load_index()
        entries.append(entry)
        save_index(entries)
    # 同步登记进 SQLite materials 表（admin 上传直通正常，文件已在库内，管理员即入库闸）
    mat, _e = CASEDB.create_material(user, {
        "id": entry["materialId"], "fileId": entry["id"],
        "title": entry["title"], "kind": "文档", "tags": ["教师上传"],
        "source": "教师上传 · " + (entry.get("byName") or entry.get("by") or ""),
        "publishedAt": entry["at"], "level": level, "credibility": "normal",
        "scope": "全体教师", "summary": entry.get("summary") or "",
    }, status="正常")
    mark_search_dirty()
    return {"ok": True, "material": mat}, 200


def find_entry(fid):
    for e in load_index():
        if e.get("id") == fid:
            return e
    return None


def _file_content_visible(entry, user):
    mid = CASEDB.material_id_for_file(entry.get("id")) if CASEDB else ""
    if mid:
        return CASEDB.can_read_material(mid, user)
    return _level_content_visible(entry.get("level"), user)


def _level_content_visible(value, user):
    try:
        level = int(value)
    except (TypeError, ValueError):
        level = 2
    if level == 0:
        return True
    return bool(user) if level == 1 else bool(user and user.get("admin"))


FILE_CTYPES = {
    ".md": "text/markdown; charset=utf-8", ".markdown": "text/markdown; charset=utf-8",
    ".txt": "text/plain; charset=utf-8", ".pdf": "application/pdf",
    ".doc": "application/msword",
    ".docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    ".xls": "application/vnd.ms-excel",
    ".xlsx": "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
    ".ppt": "application/vnd.ms-powerpoint",
    ".pptx": "application/vnd.openxmlformats-officedocument.presentationml.presentation",
    ".png": "image/png", ".jpg": "image/jpeg", ".jpeg": "image/jpeg",
    ".zip": "application/zip", ".rar": "application/vnd.rar",
}


def api_file_download(user, fid):
    e = find_entry(fid)
    if not e:
        return None, {"ok": False, "error": "文件不存在"}, 404
    level = int(e.get("level", 0))
    if not _file_content_visible(e, user):
        return None, {"ok": False, "error":
                      "该文件密级为「%s」，超出你的授权范围" % LEVEL_NAMES[level]}, 403
    full = os.path.join(FILES_DIR, e["path"])
    if not os.path.isfile(full):
        return None, {"ok": False, "error": "文件已丢失（索引在、实体不在）"}, 404
    with open(full, "rb") as f:
        data = f.read()
    ctype = FILE_CTYPES.get(os.path.splitext(e["name"])[1].lower(),
                            "application/octet-stream")
    return (data, ctype, e["name"]), None, 200


def api_file_text(user, fid):
    """在线文本：鉴权与密级检查与文件下载一致。"""
    e = find_entry(fid)
    if not e:
        return {"ok": False, "error": "文件不存在"}, 404
    level = int(e.get("level", 0))
    if not _file_content_visible(e, user):
        return {"ok": False, "error":
                "该文件密级为「%s」，超出你的授权范围" % LEVEL_NAMES[level]}, 403
    tp = e.get("textPath")
    if not tp:
        return {"ok": False, "error":
                "该文件没有可在线查看的文本（仅 md/txt/docx 上传时自动抽取）"}, 404
    full = os.path.join(FILES_DIR, tp)
    if not os.path.isfile(full):
        return {"ok": False, "error": "文本副本已丢失（索引在、实体不在）"}, 404
    with open(full, encoding="utf-8") as f:
        return {"ok": True, "text": f.read()}, 200


def api_file_delete(user, fid):
    if not user:
        return {"ok": False, "error": "未登录或登录已过期，请重新切换账号"}, 401
    if not user.get("admin"):
        return {"ok": False, "error": "仅案例管理员可删除资料文件"}, 403
    with _INDEX_LOCK:
        entries = load_index()
        e = next((x for x in entries if x.get("id") == fid), None)
        if not e:
            return {"ok": False, "error": "文件不存在"}, 404
        if e.get("seed"):
            return {"ok": False, "error": "种子素材不能删除，只能停用"}, 403
        save_index([x for x in entries if x.get("id") != fid])
    try:
        os.remove(os.path.join(FILES_DIR, e["path"]))
    except OSError:
        pass
    if e.get("textPath"):
        try:
            os.remove(os.path.join(FILES_DIR, e["textPath"]))
        except OSError:
            pass
    CASEDB.delete_material_by_file(fid)  # 联动删除素材登记行
    mark_search_dirty()
    return {"ok": True}, 200


def _set_file_level(fid, level):
    """调整文件索引密级；素材治理（/api/materials PATCH）与文件 PATCH 共用。"""
    with _INDEX_LOCK:
        entries = load_index()
        e = next((x for x in entries if x.get("id") == fid), None)
        if not e:
            return False
        e["level"] = level
        save_index(entries)
    return True


def api_file_patch(user, fid, payload):
    """调整文件密级（仅 admin）。种子条目重建后会回到 build_data.py 默认值。"""
    if not user:
        return {"ok": False, "error": "未登录或登录已过期，请重新切换账号"}, 401
    if not user.get("admin"):
        return {"ok": False, "error": "仅案例管理员可调整文件密级"}, 403
    try:
        level = int(payload.get("level"))
    except Exception:
        level = -1
    if level not in (0, 1, 2):
        return {"ok": False, "error": "密级取值无效"}, 400
    if not _set_file_level(fid, level):
        return {"ok": False, "error": "文件不存在"}, 404
    mark_search_dirty()
    return {"ok": True}, 200


# ---------------------------------------------------------------- 知识库在线导入
def load_knowledge():
    try:
        with open(KNOWLEDGE_FILE, encoding="utf-8") as f:
            data = json.load(f)
        return data if isinstance(data, list) else []
    except Exception:
        return []


def save_knowledge(sources):
    tmp = KNOWLEDGE_FILE + ".tmp"
    with open(tmp, "w", encoding="utf-8") as f:
        json.dump(sources, f, ensure_ascii=False, indent=1)
    os.replace(tmp, KNOWLEDGE_FILE)


def _md_clean(text):
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def parse_knowledge_markdown(source_id, markdown):
    """与 tools/build_data.py 教材解析同规则：# 为章、### 为节；
    产出 chapters/sections（节含 id/title/text），id 以 sourceId 为前缀保证跨来源唯一。"""
    chapters, sections = [], []
    ch, sec, buf = None, None, []

    def flush():
        nonlocal buf
        body = _md_clean("\n".join(buf))
        buf = []
        if sec is not None:
            sections.append({
                "id": "%s-%02d-%02d" % (source_id, ch["index"], sec["index"]),
                "chapterId": ch["id"],
                "chapter": ch["title"],
                "index": sec["index"],
                "title": sec["title"],
                "text": body,
                "chars": len(body),
            })
        elif ch is not None and body:
            ch["intro"] = (ch.get("intro") or "") + body + "\n"

    for ln in markdown.splitlines():
        m1 = re.match(r"^#\s+(.+)$", ln)
        m3 = re.match(r"^###\s+(.+)$", ln)
        if m1:
            flush()
            ch = {"id": "%s-c%02d" % (source_id, len(chapters) + 1),
                  "index": len(chapters) + 1, "title": m1.group(1).strip()}
            chapters.append(ch)
            sec = None
        elif m3 and ch is not None:
            flush()
            sec = {"index": sum(1 for s in sections if s["chapterId"] == ch["id"]) + 1,
                   "title": m3.group(1).strip()}
        else:
            buf.append(ln)
    flush()
    for c in chapters:
        c["intro"] = _md_clean(c.get("intro") or "")
        c["sections"] = ["%s-%02d-%02d" % (source_id, c["index"], s["index"])
                         for s in sections if s["chapterId"] == c["id"]]
    return chapters, sections


def api_knowledge_import(user, payload):
    """导入 markdown 教材为知识来源（仅 admin）；同名来源覆盖而非追加。"""
    if not user:
        return {"ok": False, "error": "未登录或登录已过期，请重新切换账号"}, 401
    if not user.get("admin"):
        return {"ok": False, "error": "仅案例管理员可导入知识库"}, 403
    name = (payload.get("name") or "").strip()
    markdown = payload.get("markdown") or ""
    if not name:
        return {"ok": False, "error": "缺少知识来源名称 name"}, 400
    if not markdown.strip():
        return {"ok": False, "error": "缺少 markdown 内容"}, 400
    with _INDEX_LOCK:
        sources = load_knowledge()
        old = next((s for s in sources if s.get("name") == name), None)
        source_id = old["sourceId"] if old else "kn-%d-%s" % (
            int(time.time()), uuid.uuid4().hex[:4])
        chapters, sections = parse_knowledge_markdown(source_id, markdown)
        entry = {
            "sourceId": source_id, "name": name,
            "source": (payload.get("source") or "").strip(),
            "importedAt": time.strftime("%Y-%m-%d %H:%M"),
            "by": user["id"], "byName": user.get("name", ""),
            "chapters": chapters, "sections": sections,
        }
        sources = [s for s in sources if s.get("name") != name] + [entry]
        save_knowledge(sources)
    mark_search_dirty()
    return {"ok": True, "sourceId": source_id, "name": name,
            "chapters": len(chapters), "sections": len(sections)}, 200


# ------------------------------------------------------------ 服务端检索索引（BM25）
_SEARCH_LOCK = threading.Lock()
_SEARCH_STATE = {"version": 0, "built_version": -1, "index": None}
BOOK_MD = os.path.join(FILES_DIR, "seed", "book", "zrbjf-2025.md")
BM25_K1, BM25_B = 1.5, 0.75


def mark_search_dirty():
    """语料变化（上传/删除/调整密级/知识导入）后置脏，下次检索时惰性重建索引。"""
    with _SEARCH_LOCK:
        _SEARCH_STATE["version"] += 1


def _tokenize(text):
    """中文按字符 bigram 分词，非中文（字母/数字）按词。"""
    tokens = []
    for part in re.findall(r"[a-z0-9]+|[㐀-䶿一-鿿]+", (text or "").lower()):
        if re.match(r"^[a-z0-9]+$", part):
            tokens.append(part)
        elif len(part) == 1:
            tokens.append(part)
        else:
            tokens.extend(part[i:i + 2] for i in range(len(part) - 1))
    return tokens


def chunk_md(text):
    """派生切片（ADR 0010，与前端 U.chunkMd / tools/build_data.py 同规则）：
    按 #{1,3} 标题行切标题树，地址 = 结构路径（如教材 2.1.1、学习资料 1.4）。"""
    lines = (text or "").split("\n")
    present = sorted({len(m.group(1))
                      for m in (re.match(r"^(#{1,3})\s+", l) for l in lines) if m})
    ranks = {h: i + 1 for i, h in enumerate(present)}
    chunks, counters, cur = [], [0, 0, 0, 0], None

    def flush():
        nonlocal cur
        if cur is None:
            return
        body = "\n".join(cur["buf"]).strip()
        if cur["h"] or body:
            chunks.append({"path": cur["path"], "level": cur["level"],
                           "h": cur["h"], "text": body})
        cur = None

    for ln in lines:
        m = re.match(r"^(#{1,3})\s+(.+)$", ln)
        if m:
            flush()
            r = ranks[len(m.group(1))]
            counters[r] += 1
            for i in range(r + 1, 4):
                counters[i] = 0
            cur = {"h": m.group(2).strip(), "buf": [], "level": r,
                   "path": ".".join(str(n) for n in counters[1:r + 1] if n > 0)}
        elif cur is not None:
            cur["buf"].append(ln)
        else:
            cur = {"h": "", "buf": [ln], "path": "0", "level": 0}
    flush()
    return chunks


def assign_file_secs(text, sections):
    """按 ADR 0010 规则计算每个知识节（### 标题）的文件结构路径，与 build_data.py 一致。"""
    lines = text.splitlines()
    levels = sorted({len(re.match(r"^(#{1,3})\s+", l).group(1))
                     for l in lines if re.match(r"^#{1,3}\s+", l)})
    ranks = {h: i + 1 for i, h in enumerate(levels)}
    counters = [0, 0, 0, 0]
    sec_i = 0
    seen_chapter = False
    out = {}
    for ln in lines:
        m = re.match(r"^(#{1,3})\s+", ln)
        if not m:
            continue
        r = ranks[len(m.group(1))]
        counters[r] += 1
        for i in range(r + 1, 4):
            counters[i] = 0
        if re.match(r"^#\s+", ln):
            seen_chapter = True
        elif re.match(r"^###\s+", ln) and seen_chapter and sec_i < len(sections):
            out[sections[sec_i]["id"]] = ".".join(str(n) for n in counters[1:r + 1] if n > 0)
            sec_i += 1
    return out


def _build_search_index():
    """收集语料并构建 Okapi BM25 索引。
    doc: {id, cls(knowledge|material|case), title, chapter, source, text, level,
          credibility, materialId?, sec?, secTitle?, status?, ownerId?}"""
    docs = []

    def add(cls, doc_id, title, chapter, text, source="", level=0, credibility="", **extra):
        text = (text or "").strip()
        if not doc_id or not text:
            return
        d = {
            "id": doc_id, "cls": cls, "title": (title or doc_id).strip(),
            "chapter": (chapter or "").strip(), "source": source, "text": text,
            "level": level, "credibility": credibility,
        }
        d.update(extra)
        docs.append(d)

    # 1) 教材：与知识库导入同规则（# 章 / ### 节）切节；节 id 与前端知识库一致（kn-xx-xx），
    #    sec 为教材文件的结构路径切片锚点（ADR 0010 深链）
    try:
        with open(BOOK_MD, encoding="utf-8") as f:
            book_text = f.read()
        _chapters, sections = parse_knowledge_markdown("kn", book_text)
        secs_map = assign_file_secs(book_text, sections)
        for s in sections:
            add("knowledge", s["id"], s["title"], s["chapter"], s["text"],
                sec=secs_map.get(s["id"], ""))
    except OSError:
        pass

    # 2) 素材（SQLite 登记库，WP2）：仅 status=正常 的进语料（入库闸，候选不进生成语料）。
    #    正文 = 标题+摘要+内容副本；上传素材有抽取文本（textPath）的按 ADR 0010 切片入索引，
    #    命中带 materialId + sec 结构路径，前端可拼 #/material/<id>?sec=<path> 深链
    if CASEDB is not None:
        entries_by_fid = {e.get("id"): e for e in load_index()}
        for m in CASEDB.materials_for_index():
            add("material", m["id"], m["title"], "",
                "\n".join([m["title"], m["summary"], m["excerpt"]]),
                source=m["source"], level=m["level"], credibility=m["credibility"],
                materialId=m["id"], grade=m["grade"], kind=m["kind"],
                publishedAt=m["publishedAt"], namePublic=bool(m.get("mountCount")))
            e = entries_by_fid.get(m["fileId"]) if m["fileId"] else None
            tp = e.get("textPath") if e else None
            if not tp:
                continue
            try:
                with open(os.path.join(FILES_DIR, tp), encoding="utf-8") as f:
                    text = f.read(30000)
            except OSError:
                continue
            for ck in chunk_md(text):
                add("material", m["id"], m["title"], "",
                    ((ck["h"] + "\n") if ck["h"] else "") + ck["text"],
                    source=m["source"], level=m["level"], credibility=m["credibility"],
                    materialId=m["id"], grade=m["grade"], kind=m["kind"],
                    publishedAt=m["publishedAt"], sec=ck["path"], secTitle=ck["h"],
                    namePublic=bool(m.get("mountCount")))

    # 3) 运行时导入的知识条目（无 knowledge.json 时跳过）
    for src in load_knowledge():
        for s in src.get("sections") or []:
            add("knowledge", s.get("id"), s.get("title"),
                "%s / %s" % (src.get("name") or "知识库", s.get("chapter") or ""),
                s.get("text"))

    # 4) 案例（SQLite 业务库）：查询时按用户身份过滤（published 全员，草稿/待审仅作者与管理员）
    if CASEDB is not None:
        for c in CASEDB.cases_for_index():
            add("case", c["id"], c["title"], "", c["text"],
                status=c["status"], ownerId=c["ownerId"])

    tf = [Counter(_tokenize(d["title"] + "\n" + d["text"])) for d in docs]
    dl = [sum(c.values()) for c in tf]
    df = Counter()
    for c in tf:
        for t in c:
            df[t] += 1
    avgdl = (sum(dl) / len(dl)) if dl else 0.0
    return {"docs": docs, "tf": tf, "dl": dl, "df": df, "avgdl": avgdl, "n": len(docs)}


def get_search_index():
    """惰性构建索引；版本戳被置脏后在下次调用时重建。"""
    with _SEARCH_LOCK:
        if (_SEARCH_STATE["index"] is not None
                and _SEARCH_STATE["built_version"] == _SEARCH_STATE["version"]):
            return _SEARCH_STATE["index"]
        idx = _build_search_index()
        _SEARCH_STATE["index"] = idx
        _SEARCH_STATE["built_version"] = _SEARCH_STATE["version"]
        return idx


def _make_snippet(text, q_tokens, width=60):
    """命中位置前后各约 width 字；未命中（经由其他词命中）取开头。"""
    flat = re.sub(r"\s+", " ", text)
    pos = -1
    for t in q_tokens:
        p = flat.lower().find(t)
        if 0 <= p and (pos < 0 or p < pos):
            pos = p
    if pos < 0:
        pos = 0
    return flat[max(0, pos - width):pos + width + 2].strip()


def search_corpus(q, kinds=None, limit=8, terms=None, user=None):
    """BM25 统一打分，分 knowledge / materials / cases 三类返回；
    materials 按用户密级过滤，cases 仅 published 对所有人可见（草稿/待审限作者与管理员）。
    terms 为前端扩展后的查询词（缺省对 q 做 bigram 分词）。"""
    empty = {"knowledge": [], "materials": [], "cases": []}
    q_tokens = [t.lower() for t in terms if t and t.strip()] if terms else _tokenize(q)
    idx = get_search_index()
    if not q_tokens or not idx["n"]:
        return empty
    if kinds:
        mapped = set()
        for k in kinds:
            k = {"materials": "material", "cases": "case"}.get(k, k)
            if k in ("knowledge", "material", "case"):
                mapped.add(k)
        kinds = mapped or None
    avgdl = idx["avgdl"] or 1.0
    scored, material_access = [], {}
    for i, d in enumerate(idx["docs"]):
        if kinds and d["cls"] not in kinds:
            continue
        readable = True
        if d["cls"] == "material":
            mid = d.get("materialId") or d["id"]
            if mid not in material_access:
                material_access[mid] = (CASEDB.can_read_material(mid, user) if CASEDB
                                        else _level_content_visible(d["level"], user))
            readable = material_access[mid]
            if not readable and not d.get("namePublic"):
                continue
        if d["cls"] == "case" and d.get("status") != "published":
            if not user or (not user.get("admin") and d.get("ownerId") != user["id"]):
                continue
        c = idx["tf"][i] if readable else Counter(_tokenize(d["title"]))
        dl = sum(c.values()) or 1
        score = 0.0
        for t in q_tokens:
            n = idx["df"].get(t, 0)
            f = c.get(t, 0)
            if not n or not f:
                continue
            idf = math.log(1 + (idx["n"] - n + 0.5) / (n + 0.5))
            score += idf * (f * (BM25_K1 + 1)) / (f + BM25_K1 * (1 - BM25_B + BM25_B * dl / avgdl))
        if score > 0:
            scored.append((score, i, readable))
    scored.sort(key=lambda x: -x[0])
    knowledge, materials, cases = [], [], []
    seen_materials = set()
    for score, i, readable in scored:
        d = idx["docs"][i]
        if d["cls"] == "knowledge" and len(knowledge) < limit:
            hit = {
                "id": d["id"], "title": d["title"], "chapter": d["chapter"],
                "snippet": _make_snippet(d["text"], q_tokens), "score": round(score, 3),
            }
            if d.get("sec"):
                hit["sec"] = d["sec"]
            knowledge.append(hit)
        elif d["cls"] == "material" and len(materials) < limit:
            if d["id"] in seen_materials:  # 一个素材多个切片命中时只保留最高分切片
                continue
            seen_materials.add(d["id"])
            hit = {
                "id": d["id"], "title": d["title"],
                "source": d["source"] if readable else "",
                "snippet": _make_snippet(d["text"], q_tokens) if readable else "",
                "score": round(score, 3), "level": d["level"],
                "credibility": d["credibility"] if readable else "",
                "grade": d.get("grade", "") if readable else "",
                "kind": d.get("kind", ""),
                "publishedAt": d.get("publishedAt", "") if readable else "",
                "materialId": d.get("materialId") or d["id"],
                "contentAvailable": readable,
            }
            if readable and d.get("sec"):
                hit["sec"] = d["sec"]
                hit["secTitle"] = d.get("secTitle", "")
            materials.append(hit)
        elif d["cls"] == "case" and len(cases) < limit:
            cases.append({
                "id": d["id"], "title": d["title"],
                "snippet": _make_snippet(d["text"], q_tokens), "score": round(score, 3),
                "status": d.get("status", ""), "ownerId": d.get("ownerId", ""),
            })
    return {"knowledge": knowledge, "materials": materials, "cases": cases}


# ------------------------------------------------------------ 引用证据（WP3）
_BOOK_SECS_CACHE = {}


def _book_sections():
    """教材节（含 fileSec 切片锚点），惰性缓存；供引用证据回填与 chunk 组装。"""
    if "secs" not in _BOOK_SECS_CACHE:
        try:
            with open(BOOK_MD, encoding="utf-8") as f:
                text = f.read()
            _c, secs = parse_knowledge_markdown("kn", text)
            fsecs = assign_file_secs(text, secs)
            for s in secs:
                s["fileSec"] = fsecs.get(s["id"], "")
            _BOOK_SECS_CACHE["secs"] = secs
        except OSError:
            _BOOK_SECS_CACHE["secs"] = []
    return _BOOK_SECS_CACHE["secs"]


def _evidence_for_target(target):
    """按引用目标 best-effort 取证据片段：kn 节 → fileSec + 节正文摘录；
    素材 → 内容副本 excerpt（无 excerpt 但有文件文本时取首个切片）。"""
    if not target:
        return None
    now = time.strftime("%Y-%m-%d %H:%M")
    if target.startswith("kn-"):
        s = next((x for x in _book_sections() if x["id"] == target), None)
        if not s:
            return None
        return {"materialId": target, "sec": s.get("fileSec") or "",
                "snippet": re.sub(r"\s+", " ", s.get("text") or "")[:200],
                "capturedAt": now}
    if CASEDB is None:
        return None
    m = CASEDB.get_material_raw(target)
    if not m:
        return None
    snippet = re.sub(r"\s+", " ", m.get("excerpt") or m.get("summary") or "")[:200]
    sec = ""
    if not snippet and m.get("fileId"):
        e = next((x for x in load_index() if x.get("id") == m["fileId"]), None)
        tp = e.get("textPath") if e else None
        if tp:
            try:
                with open(os.path.join(FILES_DIR, tp), encoding="utf-8") as f:
                    text = f.read(30000)
                chunks = chunk_md(text)
                if chunks:
                    sec = chunks[0]["path"] if len(chunks) > 1 else ""
                    snippet = re.sub(r"\s+", " ", chunks[0]["text"])[:200]
            except OSError:
                pass
    if not snippet:
        return None
    return {"materialId": m["id"], "sec": sec, "snippet": snippet, "capturedAt": now}


def api_search(user, payload):
    q = (payload.get("q") or "").strip()
    if not q:
        return {"ok": False, "error": "缺少检索词 q"}, 400
    try:
        limit = min(max(int(payload.get("limit") or 8), 1), 500)
    except Exception:
        limit = 8
    kinds = payload.get("kinds")
    if isinstance(kinds, str):
        kinds = [k.strip() for k in kinds.split(",")]
    if not isinstance(kinds, list):
        kinds = None
    terms = payload.get("terms")
    if not (isinstance(terms, list) and any(isinstance(t, str) and t.strip() for t in terms)):
        terms = None
    res = search_corpus(q, kinds=kinds, limit=limit, terms=terms, user=user)
    return {"ok": True, "q": q, "knowledge": res["knowledge"],
            "materials": res["materials"], "cases": res["cases"]}, 200


# ------------------------------------------------------------ Neo4j 图谱库（WP7）
def _graph_snapshot():
    """灌库数据集：教材骨架（章/节，kn 前缀与检索索引一致）+ SQLite 案例/素材全量。"""
    chapters, sections = [], []
    try:
        with open(BOOK_MD, encoding="utf-8") as f:
            book = f.read()
        chs, secs = parse_knowledge_markdown("kn", book)
        fsecs = assign_file_secs(book, secs)
        chapters = [{"id": c["id"], "title": c["title"], "index": c["index"]} for c in chs]
        sections = [{"id": s["id"], "title": s["title"], "chapter": s["chapter"],
                     "chapterId": s["chapterId"], "index": s["index"],
                     "fileSec": fsecs.get(s["id"], ""),
                     "snippet": re.sub(r"\s+", " ", s.get("text") or "")[:200]}
                    for s in secs]
    except OSError:
        pass
    return {"chapters": chapters, "sections": sections,
            "cases": CASEDB.cases_for_graph() if CASEDB else [],
            "materials": CASEDB.materials_for_graph() if CASEDB else []}


def _graph_enrich(gen):
    """LLM 实体增强（后台线程，可整体跳过）：逐素材从 摘要+内容副本 抽 3-6 个主题词、
    挂 0-3 个教材节（MENTIONS/HAS_TAG src=llm）。AI 未配置时跳过，图谱基础结构照样可用。"""
    if not (AI_BASE_URL and AI_API_KEY):
        return
    kn_menu = "\n".join("%s %s" % (s["id"], s["title"]) for s in _graph_snapshot()["sections"])
    if not kn_menu:
        return
    kn_ids = {ln.split(" ", 1)[0] for ln in kn_menu.splitlines()}
    for m in CASEDB.materials_for_graph():
        if gen != GRAPH.gen:
            return
        text = (m["title"] + "\n" + (m.get("summary") or "") + "\n"
                + (m.get("excerpt") or ""))[:800].strip()
        if not text:
            continue
        content, err = llm_call(AI_DEFAULT_MODEL, [
            {"role": "system", "content":
             "你是思政教学资源知识标引员。只输出 JSON：{\"kn\":[教材节id],\"topics\":[主题词]}。"
             "kn 从给定教材节目录选 0-3 个最相关的节 id；topics 提取 3-6 个关键知识点/主题词"
             "（2-8 字名词短语，不要「材料」「政策」这类泛词）。"},
            {"role": "user", "content":
             "教材节目录：\n%s\n\n素材：\n%s" % (kn_menu, text)}],
            temperature=0.2, max_tokens=300)
        obj = _extract_json(content) if not err else None
        if not isinstance(obj, dict):
            continue
        topics = [str(t).strip() for t in (obj.get("topics") or [])
                  if isinstance(t, str) and 2 <= len(t.strip()) <= 12][:6]
        GRAPH.add_mentions(m["id"], [k for k in (obj.get("kn") or []) if k in kn_ids][:3],
                           topics, gen=gen)


_GRAPH_REBUILD = {"busy": False, "again": False}


def graph_rebuild():
    """全量重建图谱（后台线程；进行中再次触发则结束后补一轮）。返回是否已启动。"""
    if not GRAPH.configured or CASEDB is None:
        return False
    if _GRAPH_REBUILD["busy"]:
        _GRAPH_REBUILD["again"] = True
        return True

    def job():
        _GRAPH_REBUILD["busy"] = True
        try:
            while True:
                _GRAPH_REBUILD["again"] = False
                counts = GRAPH.rebuild(_graph_snapshot())
                if counts is not None:
                    sys.stderr.write("[graph] 图谱重建完成：%s\n"
                                     % json.dumps(counts, ensure_ascii=False))
                    threading.Thread(target=_graph_enrich, args=(GRAPH.gen,), daemon=True).start()
                if not _GRAPH_REBUILD["again"]:
                    break
        finally:
            _GRAPH_REBUILD["busy"] = False

    threading.Thread(target=job, daemon=True).start()
    return True


def _graph_changed(kind, nid=None):
    """db.py 写路径钩子（db.GRAPH_HOOK）：案例/素材增删改同步图谱，reseed 后全量重建。"""
    if not GRAPH.configured or CASEDB is None:
        return
    try:
        if kind == "rebuild":
            graph_rebuild()
        elif kind == "case":
            c = CASEDB.case_for_graph(nid)
            if c:
                GRAPH.sync_case(c)
        elif kind == "case_deleted":
            GRAPH.delete_case(nid)
        elif kind == "material":
            m = CASEDB.material_for_graph(nid)
            if m:
                GRAPH.sync_material(m)
        elif kind == "material_deleted":
            GRAPH.delete_material(nid)
    except Exception as e:
        sys.stderr.write("[graph] 增量同步失败（%s %s）：%s\n" % (kind, nid, e))


def _graph_unavailable():
    return {"ok": False, "degraded": True,
            "error": "图谱服务不可用（Neo4j 未连接），检索等其他功能不受影响"}, 503


def _graph_filter_visible(sub, user):
    """案例按发布状态过滤；素材名称与内容权限分离。"""
    keep, locked = set(), set()
    for n in sub["nodes"]:
        if n["type"] == "case" and not (
                n.get("status") == "published"
                or (user and (user.get("admin") or n.get("ownerId") == user["id"]))):
            continue
        if n["type"] == "material" and not CASEDB.can_read_material(n["id"], user):
            if not CASEDB.material_name_public(n["id"]):
                continue
            n["contentAvailable"] = False
            locked.add(n["id"])
        keep.add(n["id"])
    sub["nodes"] = [n for n in sub["nodes"] if n["id"] in keep]
    sub["links"] = [l for l in sub["links"] if l["source"] in keep and l["target"] in keep
                    and l["source"] not in locked and l["target"] not in locked]
    if "props" in sub:
        sub["props"] = {k: v for k, v in sub["props"].items() if k in keep and k not in locked}
    return sub


def api_graph_overview(user):
    """全库轻量图：教材骨架（章/节）+ 案例 + 节级引用边（素材经 ego 按需展开）。"""
    if not GRAPH.available():
        return _graph_unavailable()
    sub = GRAPH.overview()
    if sub is None:
        return _graph_unavailable()
    sub = _graph_filter_visible(sub, user)
    return {"ok": True, "nodes": sub["nodes"], "links": sub["links"],
            "stats": GRAPH.counts()}, 200


def api_graph_ego(user, query):
    """以某节点为中心的两跳子图（graph.js 渲染数据源）。"""
    ntype = (query.get("type") or [""])[0]
    nid = (query.get("id") or [""])[0]
    if not ntype or not nid:
        return {"ok": False, "error": "缺少参数 type/id"}, 400
    if not GRAPH.available():
        return _graph_unavailable()
    sub = GRAPH.ego(ntype, nid)
    if sub is None or not sub["nodes"]:
        return {"ok": False, "error": "节点不存在或图谱不可用"}, 404
    sub = _graph_filter_visible(sub, user)
    return {"ok": True, "center": nid, "nodes": sub["nodes"], "links": sub["links"]}, 200


def api_graph_reverse(user, query):
    """知识点/标签反查：直接关联的案例与素材 + 两跳路径（知识点→案例→素材）。"""
    kn = (query.get("kn") or [""])[0]
    tag = (query.get("tag") or [""])[0]
    if not kn and not tag:
        return {"ok": False, "error": "缺少参数 kn 或 tag"}, 400
    if not GRAPH.available():
        return _graph_unavailable()
    res = GRAPH.reverse(kn=kn or None, tag=tag or None)
    if res is None:
        return {"ok": False, "error": "节点不存在或图谱不可用"}, 404
    res = _graph_filter_visible(res, user)
    res["cases"] = [n for n in res["cases"] if n["id"] in {x["id"] for x in res["nodes"]}]
    res["ok"] = True
    return res, 200


_GRAPH_TYPE_NAMES = {"chapter": "教材章", "section": "教材节", "case": "案例",
                     "material": "素材", "tag": "标签"}


def api_graph_qa(user, payload):
    """全局图谱问答：BM25 种子 → 图谱一跳子图召回证据 → LLM 综合（引用=节点 id）。
    AI 不可用/失败时降级返回召回子图统计（degraded=true，answer=null）。"""
    q = (payload.get("q") or "").strip()
    if not q:
        return {"ok": False, "error": "缺少问题 q"}, 400
    if not GRAPH.available():
        return _graph_unavailable()
    res = search_corpus(q, limit=5, user=user)
    seeds = [h["id"] for h in res["knowledge"] + res["materials"] + res["cases"]]
    sub = GRAPH.neighborhood(seeds) if seeds else {"nodes": [], "links": [], "props": {}}
    if sub is None:
        return _graph_unavailable()
    sub = _graph_filter_visible(sub, user)
    stats = {}
    for n in sub["nodes"]:
        stats[_GRAPH_TYPE_NAMES.get(n["type"], n["type"])] = \
            stats.get(_GRAPH_TYPE_NAMES.get(n["type"], n["type"]), 0) + 1
    base = {"ok": True, "q": q, "refs": sub["nodes"][:12], "stats": stats,
            "graph": {"nodes": sub["nodes"], "links": sub["links"]}}
    if not (AI_BASE_URL and AI_API_KEY):
        base.update({"degraded": True, "answer": None, "note": "AI 未配置，返回召回子图统计"})
        return base, 200
    lines = []
    for n in sub["nodes"][:40]:
        p = sub["props"].get(n["id"], {})
        body = re.sub(r"\s+", " ", (p.get("summary") or p.get("snippet") or ""))[:120]
        lines.append("〔%s〕%s｜%s｜%s"
                     % (n["id"], _GRAPH_TYPE_NAMES.get(n["type"], n["type"]), n["label"], body))
    by_id = {n["id"]: n["label"] for n in sub["nodes"]}
    edges = ["%s —%s→ %s" % (by_id.get(l["source"], l["source"]), l["rel"],
                             by_id.get(l["target"], l["target"]))
             for l in sub["links"][:60]]
    content, err = llm_call(AI_DEFAULT_MODEL, [
        {"role": "system", "content":
         "你是高校思政教学案例库的图谱问答助手，回答关于整个库的全局性问题"
         "（共同主题、思政元素分布、政策关联等）。基于召回的知识图谱子图"
         "（节点=案例/素材/教材知识/标签，边=引用/涉及/包含/标签）："
         "1. 结论先行，再分点展开；2. 引用具体节点在句末标注〔节点id〕；"
         "3. 只用子图证据，未覆盖的部分明说；4. 中文简练。"},
        {"role": "user", "content": "用户问题：%s\n\n召回子图节点（〔id〕类型｜标题｜摘要）：\n%s\n\n子图关系：\n%s"
         % (q, "\n".join(lines) or "（无）", "\n".join(edges) or "（无）")}],
        temperature=0.3, max_tokens=1200)
    if err:
        base.update({"degraded": True, "answer": None,
                     "note": "AI 调用失败（%s），返回召回子图统计" % err})
        return base, 200
    cited = set(re.findall(r"〔([^〕]+)〕", content))
    base.update({"answer": content,
                 "refs": [n for n in sub["nodes"] if n["id"] in cited] or base["refs"]})
    return base, 200


def api_admin_graph_rebuild(user):
    err = _need_admin(user)
    if err:
        return err
    if not GRAPH.configured:
        return {"ok": False, "error": "未配置 NEO4J_* 或缺少 neo4j 驱动"}, 503
    graph_rebuild()
    return {"ok": True, "started": True}, 200


# ------------------------------------------------------------ 案例业务 API（SQLite 持久层，db.py）
def load_seed_cases():
    try:
        with open(CASES_SEED_FILE, encoding="utf-8") as f:
            data = json.load(f)
        return data if isinstance(data, list) else []
    except Exception:
        return []


def load_seed_materials():
    try:
        with open(MATERIALS_SEED_FILE, encoding="utf-8") as f:
            data = json.load(f)
        return data if isinstance(data, list) else []
    except Exception:
        return []


def _load_json(path, default):
    try:
        with open(path, encoding="utf-8") as f:
            return json.load(f)
    except Exception:
        return default


# 教师主动自检的反例库与规范词库
REVIEW_LEXICON = _load_json(LEXICON_FILE, {})
REVIEW_COUNTEREXAMPLES = _load_json(COUNTEREXAMPLES_FILE, [])


def _need_login(user):
    if not user:
        return {"ok": False, "error": "未登录或登录已过期，请重新切换账号"}, 401
    return None


def api_cases_list(user, query):
    status = (query.get("status") or [""])[0].strip()
    owner = (query.get("ownerId") or [""])[0].strip()
    return {"ok": True,
            "cases": CASEDB.list_cases(user, status or None, owner or None)}, 200


def api_case_get(user, cid):
    c = CASEDB.get_case(cid, user)
    if not c:
        return {"ok": False, "error": "案例不存在或无权查看"}, 404
    return {"ok": True, "case": c}, 200


def api_case_create(user, payload):
    err = _need_login(user)
    if err:
        return err
    payload = dict(payload or {})
    payload.setdefault("id", "c-" + uuid.uuid4().hex[:10])
    c, e = CASEDB.create_case(user, payload)
    if e:
        return {"ok": False, "error": e}, 409
    mark_search_dirty()
    return {"ok": True, "case": c}, 200


def api_case_patch(user, cid, payload):
    err = _need_login(user)
    if err:
        return err
    c, e = CASEDB.update_case(user, cid, payload or {})
    if e:
        return {"ok": False, "error": e}, 404 if "不存在" in e else 403
    if "blocks" in (payload or {}):
        # 兼容编辑器改 blocks 后同步 docx（docx 为编辑真源，仅对已 docx 化的案例）
        sync_docx_from_blocks(cid, c.get("blocks"), only_if_exists=True)
    mark_search_dirty()
    return {"ok": True, "case": c}, 200


def _owned_case(user, cid):
    if not user:
        return None, ({"ok": False, "error": "未登录或登录已过期"}, 401)
    c = CASEDB.get_case(cid, user)
    if not c:
        return None, ({"ok": False, "error": "案例不存在"}, 404)
    if c.get("ownerId") != user["id"]:
        return None, ({"ok": False, "error": "仅作者本人可修改附件"}, 403)
    return c, None


def _attachment_raw(payload):
    data = payload.get("data") or ""
    if not data:
        return b"", None
    try:
        raw = base64.b64decode(data, validate=True)
    except Exception:
        return b"", "文件内容不是有效 base64"
    return (raw, None) if len(raw) <= MAX_UPLOAD_BYTES else (b"", "文件超过 20MB")


def _store_case_attachment(cid, aid, name, raw):
    if not raw:
        return ""
    ext = os.path.splitext(name)[1].lower()[:10]
    folder = os.path.join(CASE_ATTACHMENTS_DIR, cid)
    os.makedirs(folder, exist_ok=True)
    path = os.path.join(folder, aid + ext)
    with open(path, "wb") as f:
        f.write(raw)
    return os.path.relpath(path, CASE_ATTACHMENTS_DIR)


def _attachment_obj(payload, cid, raw):
    aid = "att-" + uuid.uuid4().hex[:12]
    name = (payload.get("fileName") or payload.get("title") or "附件").strip()[:180]
    text = (payload.get("text") or payload.get("excerpt")
            or payload.get("sourceUrl") or "")[:100000]
    path = _store_case_attachment(cid, aid, name, raw)
    return {"id": aid, "title": (payload.get("title") or name).strip()[:180],
            "kind": payload.get("kind") or ("网页" if payload.get("sourceUrl") else "文件"),
            "source": (payload.get("source") or "").strip()[:120],
            "sourceUrl": (payload.get("sourceUrl") or "").strip(), "fileName": name,
            "mime": payload.get("mime") or "application/octet-stream", "size": len(raw),
            "storedPath": path, "excerpt": (payload.get("excerpt") or text)[:1200],
            "contentHash": hashlib.sha256(raw or text.encode("utf-8")).hexdigest(),
            "originMaterialId": payload.get("originMaterialId") or "",
            "level": int(payload.get("level", 2)), "at": time.strftime("%Y-%m-%d %H:%M")}


def _attachment_content(payload, raw):
    if raw or not payload.get("text"):
        return payload, raw
    payload = dict(payload)
    payload["fileName"] = payload.get("fileName") or "网页原文.txt"
    payload["mime"] = "text/plain; charset=utf-8"
    return payload, payload["text"].encode("utf-8")[:MAX_UPLOAD_BYTES]


def _attachment_duplicate(c, a):
    return next((x for x in c.get("attachments") or []
                 if x.get("contentHash") == a["contentHash"]), None)


def api_case_attachment_add(user, cid, payload):
    c, err = _owned_case(user, cid)
    if err:
        return err
    payload = payload or {}
    try:
        level = int(payload.get("level", 2))
    except (TypeError, ValueError):
        level = -1
    if level not in (0, 1, 2):
        return {"ok": False, "error": "内容访问级别无效"}, 400
    payload["level"] = level
    raw, raw_err = _attachment_raw(payload)
    if raw_err:
        return {"ok": False, "error": raw_err}, 400
    payload, raw = _attachment_content(payload, raw)
    a = _attachment_obj(payload, cid, raw)
    duplicate = _attachment_duplicate(c, a)
    if duplicate:
        _remove_attachment_file(a)
        return {"ok": True, "attachment": duplicate, "case": c, "duplicate": True}, 200
    items = list(c.get("attachments") or []) + [a]
    fresh, e = CASEDB.update_case(user, cid, {"attachments": items})
    if e:
        return {"ok": False, "error": e}, 403
    mark_search_dirty()
    return {"ok": True, "attachment": a, "case": fresh}, 200


def _remove_attachment_file(a):
    path = a.get("storedPath") or ""
    full = os.path.abspath(os.path.join(CASE_ATTACHMENTS_DIR, path))
    root = os.path.abspath(CASE_ATTACHMENTS_DIR) + os.sep
    if path and full.startswith(root):
        try:
            os.remove(full)
        except OSError:
            pass


def api_case_attachment_delete(user, cid, aid):
    c, err = _owned_case(user, cid)
    if err:
        return err
    a = next((x for x in c.get("attachments") or [] if x.get("id") == aid), None)
    if not a:
        return {"ok": False, "error": "附件不存在"}, 404
    _remove_attachment_file(a)
    items = [x for x in c.get("attachments") or [] if x.get("id") != aid]
    fresh, e = CASEDB.update_case(user, cid, {"attachments": items})
    return ({"ok": True, "case": fresh}, 200) if not e else ({"ok": False, "error": e}, 403)


def _visible_case_attachment(user, cid, aid):
    c = CASEDB.get_case(cid, user)
    if not c:
        return None, None
    pools = [c.get("attachments") or []]
    if c.get("status") == "published":
        pools.append((c.get("publishedSnapshot") or {}).get("attachments") or [])
    return c, next((a for pool in pools for a in pool if a.get("id") == aid), None)


def api_case_attachment_file(user, cid, aid):
    _case, a = _visible_case_attachment(user, cid, aid)
    if a and a.get("contentAvailable") is False:
        return None, {"ok": False, "error": "无权访问附件内容"}, 403
    path = a and a.get("storedPath")
    full = os.path.abspath(os.path.join(CASE_ATTACHMENTS_DIR, path or ""))
    root = os.path.abspath(CASE_ATTACHMENTS_DIR) + os.sep
    if not a or not path or not full.startswith(root) or not os.path.isfile(full):
        return None, {"ok": False, "error": "附件文件不存在"}, 404
    with open(full, "rb") as f:
        data = f.read()
    return (data, a.get("mime") or "application/octet-stream", a.get("fileName") or "附件"), None, 200


def api_case_delete(user, cid):
    err = _need_login(user)
    if err:
        return err
    e = CASEDB.delete_case(user, cid)
    if e:
        return {"ok": False, "error": e}, 404 if "不存在" in e else 403
    delete_case_docx(cid)
    mark_search_dirty()
    return {"ok": True}, 200


def api_case_transition(user, cid, action, payload):
    """submit/withdraw（作者）与 start/approve/reject/supplement/hide/unhide（admin）。"""
    err = _need_login(user)
    if err:
        return err
    payload = payload or {}
    c, e, code = CASEDB.transition(
        user, cid, action,
        reason=(payload.get("reason") or "").strip(),
        reason_type=(payload.get("reasonType") or "").strip(),
        offline_from=(payload.get("offlineFrom") or "").strip())
    if e:
        return {"ok": False, "error": e}, code
    mark_search_dirty()
    return {"ok": True, "case": c, "reviews": CASEDB.list_reviews(50, cid)}, 200


def api_annotation_add(user, cid, payload):
    err = _need_login(user)
    if err:
        return err
    a, e = CASEDB.add_annotation(user, cid, payload or {})
    if e:
        code = 403 if e == "无权操作批注" else 404 if e == "案例不存在" else 400
        return {"ok": False, "error": e}, code
    # OO 原生批注为唯一载体：docx 已存在时注入 comment 并 bump docxVer（前端据此重建编辑器）；
    # docx 未生成（兼容模式纯块编辑）时只入侧栏，首开 OO 后由镜像兜底
    if oo_configured() and os.path.isfile(case_docx_path(cid)):
        cmid, _anchored, cerr = oo_inject_comment(
            cid, a.get("quote") or "", a.get("text") or "",
            a.get("author") or user.get("name") or user["id"])
        if cmid is not None:
            CASEDB.set_annotation_ooid(a["id"], cmid)
            CASEDB.bump_docx_ver(cid)
            a["ooId"] = str(cmid)
        else:
            sys.stderr.write("[oo] 批注注入失败 %s: %s\n" % (cid, cerr))
    c = CASEDB.get_case(cid, user)
    return {"ok": True, "annotation": a,
            "annotations": c["annotations"] if c else [a],
            "docxVer": _oo_doc_ver(cid)}, 200


def api_annotation_patch(user, aid, payload):
    err = _need_login(user)
    if err:
        return err
    a, e = CASEDB.patch_annotation(user, aid, payload or {})
    if e:
        code = 403 if e == "无权操作批注" else 404 if e == "批注不存在" else 400
        return {"ok": False, "error": e}, code
    # 「解决」同步移除 docx 中的 OO 原生批注（OO 里删除则回调镜像置 resolved，两个方向都通）
    cid = CASEDB.annotation_case_id(aid)
    if cid and oo_configured() and a.get("ooId") and a.get("status") == "resolved":
        if oo_remove_comment(cid, a["ooId"]):
            CASEDB.bump_docx_ver(cid)
    return {"ok": True, "annotation": a,
            "docxVer": _oo_doc_ver(cid) if cid else None}, 200


def api_version_add(user, cid, payload):
    err = _need_login(user)
    if err:
        return err
    v, e = CASEDB.save_version(user, cid, (payload or {}).get("label"))
    if e:
        return {"ok": False, "error": e}, 404 if "不存在" in e else 403
    save_version_docx(cid, v.get("id"))  # docx 副本到 files/cases/versions/
    c = CASEDB.get_case(cid, user)
    return {"ok": True, "version": v, "versions": c["versions"] if c else [v]}, 200


def api_version_rollback(user, cid, vid):
    err = _need_login(user)
    if err:
        return err
    c, e = CASEDB.rollback(user, cid, vid)
    if e:
        return {"ok": False, "error": e}, 404 if "不存在" in e or "快照" in e else 403
    # 回滚后重建 docx 并 bump key，强制 DS 下次按新内容重载
    sync_docx_from_blocks(cid, c.get("blocks"))
    mark_search_dirty()
    return {"ok": True, "case": c}, 200


def api_favorites(user):
    err = _need_login(user)
    if err:
        return err
    return {"ok": True, "caseIds": CASEDB.list_favorites(user),
            "materialIds": CASEDB.list_mat_favorites(user)}, 200


def api_favorite_set(user, cid, on):
    err = _need_login(user)
    if err:
        return err
    e = CASEDB.set_favorite(user, cid, on)
    if e:
        return {"ok": False, "error": e}, 404
    return {"ok": True, "caseIds": CASEDB.list_favorites(user)}, 200


def api_like_set(user, cid, on):
    err = _need_login(user)
    if err:
        return err
    res, e = CASEDB.set_like(user, cid, on)
    if e:
        return {"ok": False, "error": e}, 404
    return {"ok": True, "likes": res["likes"], "likedBy": res["likedBy"]}, 200


def api_reviews(user):
    err = _need_login(user)
    if err:
        return err
    if not user.get("admin"):
        return {"ok": False, "error": "仅案例管理员可查看审核留痕"}, 403
    return {"ok": True, "reviews": CASEDB.list_reviews(50)}, 200


def api_reseed(user):
    """管理后台「重置演示数据」：清空业务表并重灌种子。"""
    err = _need_login(user)
    if err:
        return err
    if not user.get("admin"):
        return {"ok": False, "error": "仅案例管理员可重置数据"}, 403
    n = CASEDB.reseed(load_seed_cases(), load_seed_materials())
    mark_search_dirty()
    return {"ok": True, "cases": n}, 200


def api_review_ledger(user):
    """组织资产·被退回表达台账：reject/supplement 留痕按 reasonType 聚合（admin）。"""
    err = _need_login(user)
    if err:
        return err
    if not user.get("admin"):
        return {"ok": False, "error": "仅案例管理员可查看退回台账"}, 403
    res = CASEDB.review_ledger()
    return {"ok": True, "items": res["items"], "byType": res["byType"]}, 200


# ---------------------------------------------------------------- 素材登记（SQLite，WP2）
def api_materials_list(user, query):
    """列表 + 过滤（status/kind/grade/q）；recommendFor=案例上下文推荐、
    recentCitedBy=本人最近引用（个人层）。"""
    def g(k):
        return (query.get(k) or [""])[0].strip()
    rec_for = g("recommendFor")
    if rec_for:
        err = _need_login(user)
        if err:
            return err
        ms, e = CASEDB.recommend_materials(user, rec_for)
        if e:
            return {"ok": False, "error": e}, 404
        return {"ok": True, "materials": ms}, 200
    rec_by = g("recentCitedBy")
    if rec_by:
        err = _need_login(user)
        if err:
            return err
        ms, e = CASEDB.recent_cited_materials(user, rec_by)
        if e:
            return {"ok": False, "error": e}, 403
        return {"ok": True, "materials": ms}, 200
    return {"ok": True, "materials": CASEDB.list_materials(
        user, status=g("status") or None, kind=g("kind") or None,
        grade=g("grade") or None, q=g("q") or None)}, 200


def api_material_get(user, mid):
    m = CASEDB.get_material(mid, user)
    if not m:
        return {"ok": False, "error": "素材不存在或无权查看"}, 404
    return {"ok": True, "material": m}, 200


def api_material_create(user, payload):
    """采集入库闸（ADR 0003）：必填校验 → URL 查重 → 相似度查重（top-3，force=true 跳过）。
    所有新素材一律先落「候选」，admin 在素材管理确认（改 status=正常）后才进检索语料。"""
    err = _need_login(user)
    if err:
        return err
    p = payload or {}
    if not (p.get("title") or "").strip():
        return {"ok": False, "error": "缺少素材标题 title"}, 400
    for k, name in (("grade", "信源等级"), ("gradeReason", "定级依据"),
                    ("sourceUrl", "原始链接"), ("publishedAt", "发布时间")):
        if not str(p.get(k) or "").strip():
            return {"ok": False, "error": "入库必填项缺失：%s（%s）" % (name, k)}, 400
    if p["grade"] not in ("S", "A", "B", "C"):
        return {"ok": False, "error": "信源等级取值无效（S/A/B/C）"}, 400
    dup = CASEDB.find_material_by_url(p["sourceUrl"].strip())
    if dup:
        return {"ok": False, "code": "dup",
                "error": "该链接此前已采集过：「%s」，可直接引用，无需重复入库" % dup["title"],
                "dup": {"id": dup["id"], "title": dup["title"]}}, 409
    if not p.get("force"):
        res = search_corpus(p["title"] + " " + str(p.get("summary") or "")[:200],
                            kinds=["material"], limit=3, user=user)
        # BM25 实测：真重复 ≥70，跨主题噪声 ≤12；阈值之下视为不相似，避免逢采必拦
        similar = [{"id": h["materialId"], "title": h["title"], "source": h["source"]}
                   for h in res["materials"] if h["score"] >= 25][:3]
        if similar:
            return {"ok": False, "code": "similar",
                    "error": "库中已有相似素材，建议优先复用；确认仍要采集请带 force=true",
                    "similar": similar}, 409
    m, e = CASEDB.create_material(user, p)
    if e:
        return {"ok": False, "error": e}, 409
    return {"ok": True, "material": m}, 200


def api_material_patch(user, mid, payload):
    """治理字段（密级/状态/信源等级等）限 admin；非 admin 仅可「重新采集」刷新内容副本。"""
    err = _need_login(user)
    if err:
        return err
    m, e = CASEDB.update_material(user, mid, payload or {})
    if e:
        return {"ok": False, "error": e}, 404 if "不存在" in e else 400
    if m["fileId"] and "level" in (payload or {}):
        _set_file_level(m["fileId"], m["level"])  # 保持下载强制与界面一致
    mark_search_dirty()
    return {"ok": True, "material": m}, 200


def api_materials_batch(user, payload):
    """批量治理（admin）：body {ids:[...], patch:{level/status/grade/exempt/...}}。"""
    err = _need_login(user)
    if err:
        return err
    p = payload or {}
    ids = p.get("ids")
    if not (isinstance(ids, list) and ids):
        return {"ok": False, "error": "缺少批量对象 ids"}, 400
    out, e = CASEDB.batch_update_materials(user, ids, p.get("patch") or {})
    if e:
        return {"ok": False, "error": e}, 403
    if "level" in (p.get("patch") or {}):
        for m in out:
            if m["fileId"]:
                _set_file_level(m["fileId"], m["level"])
    mark_search_dirty()
    return {"ok": True, "updated": len(out), "materials": out}, 200


def api_mat_favorite(user, mid, on):
    err = _need_login(user)
    if err:
        return err
    e = CASEDB.set_mat_favorite(user, mid, on)
    if e:
        return {"ok": False, "error": e}, 404
    return {"ok": True, "materialIds": CASEDB.list_mat_favorites(user)}, 200


def api_materials_healthcheck(user, payload):
    """来源健康检查（admin）：对有 sourceUrl 的素材发 HEAD（失败回退 GET，超时 5s，并发 8），
    失败的标 status=来源失效；body 可选 ids 限定范围。返回结果汇总。"""
    err = _need_login(user)
    if err:
        return err
    if not user.get("admin"):
        return {"ok": False, "error": "仅案例管理员可执行来源健康检查"}, 403
    ids = (payload or {}).get("ids")
    targets = [m for m in CASEDB.list_materials(user)
               if m["sourceUrl"] and (not ids or m["id"] in ids)]

    def check(m):
        url = m["sourceUrl"]
        if not re.match(r"^https?://", url):
            return m, "非 http/https 链接"
        last = "未知错误"
        for method in ("HEAD", "GET"):  # 部分站点不支持 HEAD，用 GET 复核
            try:
                req = urllib.request.Request(url, method=method, headers={"User-Agent": UA})
                with urllib.request.urlopen(req, timeout=5,
                                            context=ssl.create_default_context()) as resp:
                    if resp.status < 400:
                        return m, None
                    last = "HTTP %s" % resp.status
            except urllib.error.HTTPError as e:
                last = "HTTP %s" % e.code
            except Exception as e:
                last = str(e)[:120]
        return m, last

    failed = []
    with ThreadPoolExecutor(max_workers=8) as ex:
        for m, err2 in ex.map(check, targets):
            if err2:
                failed.append({"id": m["id"], "title": m["title"],
                               "url": m["sourceUrl"], "error": err2})
    marked = CASEDB.mark_materials_failed([f["id"] for f in failed])
    if marked:
        mark_search_dirty()
    return {"ok": True, "checked": len(targets), "failed": failed, "marked": marked}, 200


# ---------------------------------------------------------------- AI 代理
def ai_chat(payload):
    """转发 OpenAI 兼容 chat 请求到真实模型服务。密钥不离开服务端。"""
    if not AI_BASE_URL or not AI_API_KEY:
        return {"ok": False, "error": "服务端未配置 AI_BASE_URL / AI_API_KEY"}
    body = {
        "model": payload.get("model") or AI_DEFAULT_MODEL,
        "messages": payload.get("messages") or [],
        "temperature": payload.get("temperature", 0.7),
        "stream": False,
    }
    # 模型只允许 .env 白名单内的取值
    if AI_MODELS and body["model"] not in AI_MODELS:
        body["model"] = AI_DEFAULT_MODEL
    if payload.get("max_tokens"):
        body["max_tokens"] = int(payload["max_tokens"])
    req = urllib.request.Request(
        AI_BASE_URL + "/chat/completions",
        data=json.dumps(body).encode("utf-8"),
        headers={
            "Content-Type": "application/json",
            "Authorization": "Bearer " + AI_API_KEY,
        },
        method="POST",
    )
    t0 = time.time()
    try:
        with urllib.request.urlopen(req, timeout=AI_TIMEOUT) as resp:
            data = json.loads(resp.read().decode("utf-8"))
    except urllib.error.HTTPError as e:
        detail = e.read().decode("utf-8", "ignore")[:500]
        return {"ok": False, "error": "模型服务 HTTP %s: %s" % (e.code, detail)}
    except Exception as e:  # 超时/网络错误
        return {"ok": False, "error": "模型调用失败: %s" % e}
    choice = (data.get("choices") or [{}])[0]
    msg = choice.get("message") or {}
    return {
        "ok": True,
        "content": msg.get("content", ""),
        "model": data.get("model", body["model"]),
        "usage": data.get("usage", {}),
        "elapsed_ms": int((time.time() - t0) * 1000),
    }


# ------------------------------------------------------------ URL 临时抓取
def _is_public_host(host):
    """基础 SSRF 防护：仅允许公网 http/https。"""
    try:
        infos = socket.getaddrinfo(host, None)
    except Exception:
        return False
    for info in infos:
        ip = info[4][0]
        if ip.startswith(("127.", "10.", "0.", "169.254.", "192.168.", "::1")):
            return False
        if ip.startswith("172."):
            try:
                second = int(ip.split(".")[1])
                if 16 <= second <= 31:
                    return False
            except Exception:
                return False
    return True


def _download_url(url):
    req = urllib.request.Request(url, headers={"User-Agent": UA})
    ctx = ssl.create_default_context()
    with urllib.request.urlopen(req, timeout=20, context=ctx) as resp:
        raw = resp.read(MAX_UPLOAD_BYTES + 1)
        return resp.geturl(), resp.headers.get("Content-Type", ""), raw


def _snapshot_text(raw):
    html = raw.decode("utf-8", "ignore")
    m = re.search(r"<title[^>]*>(.*?)</title>", html, re.S | re.I)
    title = re.sub(r"\s+", " ", m.group(1)).strip() if m else ""
    text = re.sub(r"<script[\s\S]*?</script>", " ", html, flags=re.I)
    text = re.sub(r"<style[\s\S]*?</style>", " ", text, flags=re.I)
    text = re.sub(r"<[^>]+>", "\n", text)
    text = re.sub(r"&nbsp;", " ", text)
    text = re.sub(r"&amp;", "&", text)
    text = re.sub(r"&lt;|&gt;|&quot;|&#39;", " ", text)
    lines = [ln.strip() for ln in text.splitlines()]
    lines = [ln for ln in lines if len(ln) >= 8]
    return title, "\n".join(lines)[:100000]


def fetch_url(payload):
    url = (payload.get("url") or "").strip()
    host = urllib.parse.urlparse(url).hostname or ""
    if not re.match(r"^https?://", url) or not _is_public_host(host):
        return {"ok": False, "error": "仅支持公网 http/https 链接"}
    try:
        final_url, ctype, raw = _download_url(url)
    except Exception as e:
        return {"ok": False, "error": "抓取失败: %s" % e, "url": url}
    if len(raw) > MAX_UPLOAD_BYTES:
        return {"ok": False, "error": "来源文件超过 20MB", "url": url}
    if not any(x in ctype for x in ("text", "html", "json")):
        name = os.path.basename(urllib.parse.urlparse(final_url).path) or "来源文件"
        return {"ok": True, "url": url, "finalUrl": final_url, "title": name,
                "data": base64.b64encode(raw).decode("ascii"), "fileName": name,
                "contentType": ctype or "application/octet-stream"}
    title, text = _snapshot_text(raw)
    return {"ok": True, "url": url, "finalUrl": final_url, "title": title,
            "text": text, "contentType": ctype}


# --------------------------------------------------------------- 联网检索
def _search_request(url, body, headers):
    req = urllib.request.Request(url, data=json.dumps(body).encode("utf-8"),
                                 headers=headers, method="POST")
    ctx = ssl.create_default_context()
    with urllib.request.urlopen(req, timeout=30, context=ctx) as resp:
        return json.loads(resp.read().decode("utf-8"))


def _anysearch_results(text):
    blocks = re.split(r"(?m)^### \d+\. ", text or "")[1:]
    results = []
    for rank, block in enumerate(blocks, 1):
        title, _, body = block.partition("\n")
        match = re.search(r"(?m)^- \*\*URL\*\*: (\S+)", body)
        if not match:
            continue
        content = re.sub(r"(?m)^- \*\*URL\*\*: \S+\s*", "", body).strip()
        results.append({"title": title.strip(), "url": match.group(1),
                        "content": content[:1200], "score": 1 / rank})
    return results


def _anysearch(query, limit):
    body = {"jsonrpc": "2.0", "id": 1, "method": "tools/call",
            "params": {"name": "search", "arguments": {
                "query": query, "max_results": limit}}}
    headers = {"Content-Type": "application/json"}
    if ANYSEARCH_API_KEY:
        headers["Authorization"] = "Bearer " + ANYSEARCH_API_KEY
    data = _search_request("https://api.anysearch.com/mcp", body, headers)
    content = ((data.get("result") or {}).get("content") or [])
    text = next((x.get("text", "") for x in content if x.get("type") == "text"), "")
    return _anysearch_results(text)


def _tavily(query, limit):
    body = {
        "api_key": TAVILY_API_KEY,
        "query": query,
        "max_results": limit,
        "include_answer": False,
        "include_raw_content": False,
    }
    data = _search_request("https://api.tavily.com/search", body,
                           {"Content-Type": "application/json"})
    return [{"title": r.get("title", ""), "url": r.get("url", ""),
             "content": (r.get("content") or "")[:1200], "score": r.get("score", 0)}
            for r in data.get("results") or []]


def web_search(payload):
    """检索公开网络资源，供教师选择并采集为案例附件。"""
    query = (payload.get("query") or "").strip()
    if not query:
        return {"ok": False, "error": "缺少检索词"}
    limit = min(int(payload.get("max_results") or 6), 10)
    if not (ANYSEARCH_API_KEY or TAVILY_API_KEY):
        return {"ok": False, "error": "服务端未配置联网检索"}
    try:
        results = _anysearch(query, limit) if ANYSEARCH_API_KEY else _tavily(query, limit)
    except urllib.error.HTTPError as e:
        detail = e.read().decode("utf-8", "ignore")[:300]
        return {"ok": False, "error": "检索服务 HTTP %s: %s" % (e.code, detail)}
    except Exception as e:
        return {"ok": False, "error": "检索失败: %s" % e}
    return {"ok": True, "query": query, "results": results}


# ------------------------------------------------------------ 自动盯源（WP5）
# 管道：抓栏目页 → AI 抽取候选条目（标题+链接+日期+摘要，含关键词过滤；
# AI 不可用时跳过该源并记录）→ URL+标题指纹双重去重 → 写 watch_items「待审」。
# 版权默认策略：只存标题+摘要+原文链接+元数据，不抓全文入库。
WATCH_INTERVAL_SECONDS = 3600
_WATCH_LOCK = threading.Lock()

WATCH_SOURCES_PRESET = [
    {"name": "新华网·时政频道", "url": "http://www.xinhuanet.com/politics/",
     "keywords": ["思政", "教育", "科技", "创新"]},
    {"name": "人民日报·评论", "url": "http://opinion.people.com.cn/",
     "keywords": ["教育", "青年", "科技"]},
    {"name": "求是网", "url": "http://www.qstheory.cn/",
     "keywords": ["理论", "教育", "文化"]},
]

WATCH_EXTRACT_PROMPT = (
    "你是思政素材采集助手。用户给你一个栏目页的链接清单（JSON 数组，含 title/url）"
    "和一组关注关键词。请从中挑出与关键词相关、适合作为高校思政教学素材的报道条目，"
    "只输出严格 JSON 数组（不要输出任何其他内容），最多 10 条："
    "[{\"title\":\"报道标题\",\"url\":\"原文链接\",\"date\":\"YYYY-MM-DD 或空字符串\","
    "\"summary\":\"一句话摘要\"}]。"
    "规则：只能从给定清单中选取，不得编造或改写 url；导航/广告/栏目名等非报道链接不要选；"
    "没有相关条目时输出 []。"
)


def _watch_fetch(url):
    """盯源抓取：栏目页 URL 由 admin 配置（可信），允许内网/本机演示源，
    不走 /api/fetch-url 的公网限制。返回 (html, error)。"""
    if not re.match(r"^https?://", url):
        return None, "仅支持 http/https 链接"
    req = urllib.request.Request(url, headers={"User-Agent": UA})
    try:
        with urllib.request.urlopen(req, timeout=15,
                                    context=ssl.create_default_context()) as resp:
            ctype = resp.headers.get("Content-Type", "")
            raw = resp.read(1024 * 1024)
    except Exception as e:
        return None, "抓取失败: %s" % e
    m = re.search(r"charset=([\w-]+)", ctype, re.I)
    enc = m.group(1) if m else "utf-8"
    try:
        return raw.decode(enc, "ignore"), None
    except Exception:
        return raw.decode("utf-8", "ignore"), None


def _watch_links(html, base_url):
    """从栏目页 HTML 提取锚链接清单（锚文本≥8字），喂给 AI 做条目抽取。"""
    links, seen = [], set()
    for m in re.finditer(r"<a\s[^>]*?href=[\"']([^\"']+)[\"'][^>]*>([\s\S]*?)</a>", html, re.I):
        text = re.sub(r"<[^>]+>", "", m.group(2))
        text = re.sub(r"\s+", " ", text).strip()
        if len(text) < 8:
            continue
        url = urllib.parse.urljoin(base_url, m.group(1))
        if not re.match(r"^https?://", url) or url in seen:
            continue
        seen.add(url)
        links.append({"title": text[:80], "url": url})
        if len(links) >= 120:
            break
    return links


def _watch_extract(src, html):
    """AI 条目抽取 + 关键词过滤。返回 (items, error)；AI 不可用返回 (None, 说明)。"""
    links = _watch_links(html, src["url"])
    if not links:
        return [], None
    if not (AI_BASE_URL and AI_API_KEY):
        return None, "AI 不可用（未配置模型），跳过该源"
    kw = "、".join(src.get("keywords") or []) or "（不限）"
    content, err = llm_call(AI_DEFAULT_MODEL, [
        {"role": "system", "content": WATCH_EXTRACT_PROMPT},
        {"role": "user", "content": "关注关键词：%s\n\n栏目页链接清单：\n%s"
         % (kw, json.dumps(links, ensure_ascii=False)[:6000])},
    ], temperature=0.1)
    if err:
        return None, "AI 抽取失败：" + err
    arr = _extract_json(content, want_array=True)
    items = []
    if isinstance(arr, list):
        for it in arr:
            if not isinstance(it, dict):
                continue
            title = str(it.get("title") or "").strip()
            url = str(it.get("url") or "").strip()
            if not title or not re.match(r"^https?://", url):
                continue
            items.append({"title": title[:120], "url": url,
                          "summary": str(it.get("summary") or "")[:300],
                          "publishedAt": str(it.get("date") or "")[:10]})
    return items[:10], None


def run_watch(source_ids=None):
    """盯源主管道（手动触发与定时线程共用）：逐源抓取抽取、去重落库、记录运行结果。"""
    if CASEDB is None:
        return {"ok": False, "error": "业务库未初始化"}
    targets = [s for s in CASEDB.list_watch_sources()
               if s["enabled"] and (not source_ids or s["id"] in source_ids)]
    results = []
    with _WATCH_LOCK:
        for s in targets:
            res = {"sourceId": s["id"], "name": s["name"], "ok": False, "added": 0}
            html, err = _watch_fetch(s["url"])
            if err:
                res["note"] = "源不可达：" + err
            else:
                items, err = _watch_extract(s, html)
                if err:
                    res["note"] = err  # AI 不可用：跳过该源并记录
                else:
                    res["ok"] = True
                    res["added"] = CASEDB.add_watch_items(s["id"], items)
                    res["note"] = ("本次无新增" if not res["added"]
                                   else "新增 %d 条候选" % res["added"])
            CASEDB.mark_watch_run(s["id"], res["added"])
            results.append(res)
    return {"ok": True, "added": sum(r["added"] for r in results), "results": results}


def watch_scheduler():
    """轻量定时扫描：每小时跑一遍 enabled 源（线程循环，不引入调度框架）。"""
    while True:
        time.sleep(WATCH_INTERVAL_SECONDS)
        try:
            res = run_watch()
            sys.stderr.write("[watch] 定时扫描完成：新增 %d 条候选\n" % res.get("added", 0))
        except Exception as e:
            sys.stderr.write("[watch] 定时扫描异常: %s\n" % e)


def _need_admin(user):
    err = _need_login(user)
    if err:
        return err
    if not user.get("admin"):
        return {"ok": False, "error": "仅案例管理员可管理盯源"}, 403
    return None


def api_watch_sources(user):
    err = _need_admin(user)
    if err:
        return err
    return {"ok": True, "sources": CASEDB.list_watch_sources()}, 200


def api_watch_source_create(user, payload):
    err = _need_admin(user)
    if err:
        return err
    s, e = CASEDB.create_watch_source(payload or {})
    if e:
        return {"ok": False, "error": e}, 400
    return {"ok": True, "source": s}, 200


def api_watch_source_patch(user, sid, payload):
    err = _need_admin(user)
    if err:
        return err
    s, e = CASEDB.update_watch_source(sid, payload or {})
    if e:
        return {"ok": False, "error": e}, 404 if "不存在" in e else 400
    return {"ok": True, "source": s}, 200


def api_watch_source_delete(user, sid):
    err = _need_admin(user)
    if err:
        return err
    if not CASEDB.delete_watch_source(sid):
        return {"ok": False, "error": "盯源不存在"}, 404
    return {"ok": True}, 200


def api_watch_items(user, query):
    err = _need_admin(user)
    if err:
        return err
    status = (query.get("status") or [""])[0].strip()
    return {"ok": True, "items": CASEDB.list_watch_items(status or None),
            "sources": CASEDB.list_watch_sources()}, 200


def api_watch_run(user, payload):
    """手动触发扫描（admin）；body 可选 sourceId 限定单个源。同步执行并返回每源结果。"""
    err = _need_admin(user)
    if err:
        return err
    sid = ((payload or {}).get("sourceId") or "").strip()
    return run_watch([sid] if sid else None), 200


def _title_sim(a, b):
    return difflib.SequenceMatcher(
        None, re.sub(r"\s+", "", a or ""), re.sub(r"\s+", "", b or "")).ratio()


def api_watch_item_import(user, iid, payload):
    """候选卡入库（admin）：走与 POST /api/materials 同一入库闸（URL 查重 + 必填 grade），
    落素材「候选」待确认定级；同事件其他待审报道作为「多方验证」附注写入摘要。"""
    err = _need_admin(user)
    if err:
        return err
    it = CASEDB.get_watch_item(iid)
    if not it:
        return {"ok": False, "error": "候选卡不存在"}, 404
    if it["status"] == "已入库":
        return {"ok": False, "error": "该候选卡已入库"}, 409
    p = payload or {}
    grade = (p.get("grade") or "B").strip()
    if grade not in ("S", "A", "B", "C"):
        return {"ok": False, "error": "信源等级取值无效（S/A/B/C）"}, 400
    dup = CASEDB.find_material_by_url(it["url"])
    if dup:  # 入库闸 URL 查重：已有人采过同一链接，直接关联不再重复入库
        CASEDB.set_watch_item(iid, "已入库", dup["id"])
        return {"ok": True, "material": dup, "note": "链接已在素材库中，已关联既有素材"}, 200
    src_name = next((s["name"] for s in CASEDB.list_watch_sources()
                     if s["id"] == it["sourceId"]), "盯源")
    peers = [x for x in CASEDB.list_watch_items("待审")
             if x["id"] != iid and _title_sim(x["title"], it["title"]) >= 0.5]
    src_names = {s["id"]: s["name"] for s in CASEDB.list_watch_sources()}
    summary = it["summary"]
    if peers:
        summary += ("\n多方验证：" + "；".join(
            "%s《%s》" % (src_names.get(x["sourceId"], "盯源"), x["title"])
            for x in peers[:3]))
    m, e = CASEDB.create_material(user, {
        "title": it["title"], "kind": "报道", "source": src_name,
        "sourceUrl": it["url"], "publishedAt": it["publishedAt"] or time.strftime("%Y-%m-%d"),
        "grade": grade,
        "gradeReason": (p.get("gradeReason") or "").strip()
                       or "盯源候选入库，管理员定级 %s 级" % grade,
        "summary": summary, "credibility": "normal",
    }, status="候选")
    if e:
        return {"ok": False, "error": e}, 409
    CASEDB.set_watch_item(iid, "已入库", m["id"])
    mark_search_dirty()
    return {"ok": True, "material": m}, 200


def api_watch_item_patch(user, iid, payload):
    """候选卡治理（admin）：目前仅「忽略」（status=已忽略，不再出现在待审列表）。"""
    err = _need_admin(user)
    if err:
        return err
    status = ((payload or {}).get("status") or "").strip()
    if status != "已忽略":
        return {"ok": False, "error": "仅支持忽略（status=已忽略）；入库请用 import"}, 400
    it, e = CASEDB.set_watch_item(iid, status)
    if e:
        return {"ok": False, "error": e}, 404
    return {"ok": True, "item": it}, 200


# ------------------------------------------------------------ 众筹贡献（WP5）
def api_contributions(user):
    err = _need_login(user)
    if err:
        return err
    return {"ok": True, "contributions": CASEDB.list_contributions(user)}, 200


def api_contribution_create(user, payload):
    """教师贡献（先审后发）：link=素材链接（URL+摘要+grade 建议）；
    kn_link=知识点-素材关联。完整案例贡献走既有「提交审核」流程，不在此提交。"""
    err = _need_login(user)
    if err:
        return err
    p = (payload or {}).get("payload") or {}
    kind = (payload or {}).get("kind") or ""
    if kind == "link":
        url = (p.get("url") or "").strip()
        if not re.match(r"^https?://", url):
            return {"ok": False, "error": "请填写有效的 http/https 链接"}, 400
        if not (p.get("title") or "").strip():
            return {"ok": False, "error": "请填写素材标题"}, 400
        if p.get("grade") and p["grade"] not in ("S", "A", "B", "C"):
            return {"ok": False, "error": "信源等级取值无效（S/A/B/C）"}, 400
        p["url"] = url
        dup = CASEDB.find_material_by_url(url)
        if dup:
            return {"ok": False, "code": "dup",
                    "error": "该链接此前已采集过：「%s」，可直接引用，无需重复贡献" % dup["title"]}, 409
    elif kind == "kn_link":
        kn_id = (p.get("knId") or "").strip()
        if not any(s["id"] == kn_id for s in _book_sections()):
            return {"ok": False, "error": "知识点不存在"}, 400
        if not CASEDB.get_material_raw((p.get("materialId") or "").strip()):
            return {"ok": False, "error": "素材不存在"}, 400
    else:
        return {"ok": False, "error": "贡献类型取值无效（link/kn_link）；"
                                      "完整案例请走案例「提交审核」流程"}, 400
    c, e = CASEDB.create_contribution(user, kind, p)
    if e:
        return {"ok": False, "error": e}, 400
    return {"ok": True, "contribution": c}, 200


def api_contribution_review(user, cid, payload):
    """admin 审核贡献：action=approve/reject。link 通过走入库闸落素材「候选」。"""
    err = _need_login(user)
    if err:
        return err
    p = payload or {}
    res, e, code = CASEDB.review_contribution(
        user, cid, (p.get("action") or "").strip(), (p.get("reason") or "").strip())
    if e:
        return {"ok": False, "error": e}, code
    contrib, material = res
    out = {"ok": True, "contribution": contrib}
    if material:
        out["material"] = material
        mark_search_dirty()
    return out, 200


def api_my_impact(user):
    """作者「被引用/被改编」聚合：素材贡献被引次数 + 案例被收藏/被点赞数。"""
    err = _need_login(user)
    if err:
        return err
    return dict({"ok": True}, **CASEDB.my_impact(user)), 200


# ------------------------------------------------------------ 教师生成偏好（WP4b）
def api_my_prefs_get(user):
    err = _need_login(user)
    if err:
        return err
    return {"ok": True, "prefs": CASEDB.get_prefs(user["id"])}, 200


def api_my_prefs_put(user, payload):
    """整体覆盖：{length, style, bannedWords, themes}；四项全空即清空。"""
    err = _need_login(user)
    if err:
        return err
    return {"ok": True, "prefs": CASEDB.set_prefs(user["id"], payload or {})}, 200


PREF_LABELS = (("length", "篇幅"), ("style", "语言风格"),
               ("bannedWords", "禁用词（生成结果中绝对不得出现）"), ("themes", "常结合的思政主题"))


def prefs_prompt_block(prefs):
    """教师偏好注入 prompt 的说明段；无偏好（全空）返回空串。"""
    lines = ["%s：%s" % (label, (prefs or {}).get(k, ""))
             for k, label in PREF_LABELS if (prefs or {}).get(k)]
    if not lines:
        return ""
    return ("【教师偏好（教师本人填写，生成时必须遵守）】\n" + "\n".join(lines))


def banned_word_hits(text, prefs):
    """禁用词服务端检查：逗号/顿号分隔，命中词按序去重返回。"""
    words = [w.strip() for w in re.split(r"[,，、;；]", (prefs or {}).get("bannedWords") or "")]
    return [w for w in dict.fromkeys(w for w in words if w) if w in (text or "")]


# ------------------------------------------------------------ MoA 多智能体编排（/api/ai/agent）
INTENT_ROUTES = {
    "find-theory": "librarian", "find-material": "librarian",
    "draft": "writer", "polish": "writer", "adapt": "writer", "section-draft": "writer",
    "review": "reviewer",
}

ORCHESTRATOR_PROMPT = (
    "你是思政教学案例平台的主 Agent，只负责判断用户请求应交给哪个角色处理。"
    "只输出严格 JSON（不要输出任何其他内容）："
    "{\"route\":\"librarian|writer|reviewer\",\"reason\":\"一句话理由\"}。"
    "路由规则：查找理论依据、查找/关联素材、整理来源出处 → librarian；"
    "起草、润色、改写、续写章节等文本生成 → writer；"
    "审校、核查理论准确性/事实与引用/语言规范/风险 → reviewer。"
)

LIBRARIAN_PROMPT = (
    "你是思政教学案例平台的「资料管理员」，擅长素材关联和来源整理。"
    "你可以使用三个工具：search_corpus（检索平台案例、知识和案例附件）、"
    "web_search（检索互联网公开资源）和 fetch_url（抓取指定公开网页）。"
    "每一轮只输出一个严格 JSON，不要输出任何其他文字：\n"
    "{\"action\":\"search_corpus\",\"args\":{\"q\":\"检索词\"}}\n"
    "或 {\"action\":\"web_search\",\"args\":{\"q\":\"检索词\"}}\n"
    "或 {\"action\":\"fetch_url\",\"args\":{\"url\":\"https://...\"}}\n"
    "或 {\"action\":\"final\",\"content\":\"最终答复\"}\n"
    "规则：平台资料用 search_corpus；用户明确要求联网时用 web_search；给出具体 URL 时用 fetch_url；"
    "可换检索词多次检索，信息足够后立即输出 final。"
    "final 的 content 用中文，分条列出推荐内容（标题/章节/来源），引用教材或素材用〔n〕编号标注，"
    "不要编造语料中不存在的条目。"
)

WRITER_PROMPT = (
    "你是思政教学案例平台的「写作手」，擅长教学案例文本的起草、润色与改写。"
    "根据用户意图和案例上下文直接输出成稿文本本身，不要解释、不要加任何前后缀。"
    "保持思政教学语言规范、理论表述准确。"
    "文本分三层，写法不同："
    "1) 事实陈述（人物、事件、数据、文件、时间等）必须依据用户消息中给出的资料 chunk，"
    "句末用〔n〕标注对应 chunk 编号，不得新增任何无来源事实；"
    "2) 理论解释依据知识库（教材）chunk 表述，同样用〔n〕标注；"
    "3) 教学设计类内容（目标、流程、讨论题等）可以发挥，但不得新增事实。"
    "用户消息里已关联的既有引用保持其原编号；新增引用只能使用本次检索到的 chunk 编号。"
    "资料不足以支撑时必须输出[资料不足]并说明缺什么，严禁按常识补齐事实。"
)

REVIEWER_PROMPT = (
    "你是思政教学案例平台的「内容审校员」。按标准清单逐条审校用户给出的文本："
    "理论准确性、事实与引用一致性、语言规范、风险与待人工确认。"
    "另外对用户消息中每个带〔n〕引用的句子做引用蕴含判断：对应 chunk 的内容是否真实支持该句；"
    "不支持（unsupported）时输出一条 standard 为「引用蕴含」、status 为 risk 的条目，"
    "note 说明 chunk 与实际句子的出入，ref 填被引用句原文；支持的不必逐条报告。"
    "只输出严格 JSON 数组，不要输出任何其他内容："
    "[{\"standard\":\"标准名\",\"status\":\"pass|risk|confirm\",\"note\":\"一句话说明\",\"ref\":\"〔n〕可选\"}]。"
    "status 含义：pass=通过，risk=有风险需修改，confirm=需人工确认。"
)

MOA_AGENT_META = {
    "orchestrator": {"label": "主Agent", "skill": "意图识别与任务分派"},
    "librarian": {"label": "资料管理员", "skill": "素材关联和来源整理"},
    "writer": {"label": "写作手", "skill": "双候选写作（起草/润色/改写）"},
    "reviewer": {"label": "内容审校员", "skill": "理论/事实/语言/风险逐项审校"},
}


def llm_call(model, messages, temperature=0.7, max_tokens=None):
    """非流式 LLM 调用（服务端内部编排用；模型来自服务端配置，不走客户端白名单）。
    返回 (content, error)。"""
    body = {"model": model, "messages": messages,
            "temperature": temperature, "stream": False}
    if max_tokens:
        body["max_tokens"] = int(max_tokens)
    req = urllib.request.Request(
        AI_BASE_URL + "/chat/completions",
        data=json.dumps(body).encode("utf-8"),
        headers={"Content-Type": "application/json",
                 "Authorization": "Bearer " + AI_API_KEY},
        method="POST",
    )
    try:
        with urllib.request.urlopen(req, timeout=AI_TIMEOUT) as resp:
            data = json.loads(resp.read().decode("utf-8"))
    except urllib.error.HTTPError as e:
        detail = e.read().decode("utf-8", "ignore")[:300]
        return None, "模型服务 HTTP %s: %s" % (e.code, detail)
    except Exception as e:
        return None, "模型调用失败: %s" % e
    msg = ((data.get("choices") or [{}])[0]).get("message") or {}
    return msg.get("content") or "", None


def llm_stream_call(model, messages, on_delta, temperature=0.7, max_tokens=None):
    """流式 LLM 调用：每个增量文本回调 on_delta。返回 (full_text, error)。"""
    body = {"model": model, "messages": messages,
            "temperature": temperature, "stream": True}
    if max_tokens:
        body["max_tokens"] = int(max_tokens)
    req = urllib.request.Request(
        AI_BASE_URL + "/chat/completions",
        data=json.dumps(body).encode("utf-8"),
        headers={"Content-Type": "application/json",
                 "Authorization": "Bearer " + AI_API_KEY},
        method="POST",
    )
    try:
        resp = urllib.request.urlopen(req, timeout=AI_TIMEOUT)
    except urllib.error.HTTPError as e:
        detail = e.read().decode("utf-8", "ignore")[:300]
        return None, "模型服务 HTTP %s: %s" % (e.code, detail)
    except Exception as e:
        return None, "模型调用失败: %s" % e
    full = []
    try:
        for raw_line in resp:
            line = raw_line.decode("utf-8", "ignore").strip()
            if not line.startswith("data:"):
                continue
            data = line[5:].strip()
            if data == "[DONE]":
                break
            try:
                obj = json.loads(data)
            except Exception:
                continue
            delta = ((obj.get("choices") or [{}])[0]).get("delta") or {}
            piece = delta.get("content") or ""
            if piece:
                full.append(piece)
                try:
                    on_delta(piece)
                except Exception:
                    pass
    except Exception as e:
        return ("".join(full) or None), "流式传输中断: %s" % e
    finally:
        resp.close()
    return "".join(full), None


def _extract_json(text, want_array=False):
    """从模型输出中提取第一个 JSON 对象/数组并解析；失败返回 None。"""
    if not text:
        return None
    text = text.strip()
    m = re.search(r"```(?:json)?\s*([\s\S]*?)```", text)
    if m:
        text = m.group(1).strip()
    start = text.find("[" if want_array else "{")
    if start < 0:
        return None
    try:
        obj, _end = json.JSONDecoder().raw_decode(text[start:])
        return obj
    except Exception:
        return None


def _chunk_text(text, size=24):
    for i in range(0, len(text), size):
        yield text[i:i + size]


def _moa_search_lines(res):
    lines = ["〔c%d〕案例《%s》：%s" % (i, x["title"], x["snippet"])
             for i, x in enumerate(res["cases"], 1)]
    lines += ["〔k%d〕%s / %s：%s" % (i, x["chapter"] or "知识条目", x["title"], x["snippet"])
              for i, x in enumerate(res["knowledge"], 1)]
    lines += ["〔m%d〕素材《%s》（%s）：%s" % (i, x["title"], x["source"], x["snippet"])
              for i, x in enumerate(res["materials"], 1)]
    return lines


def _moa_search_items(res):
    items = [{"kind": "case", "id": x["id"], "title": x["title"],
              "excerpt": x["snippet"], "source": "已发布案例"} for x in res["cases"]]
    items += [{"kind": "knowledge", "id": x["id"], "title": x["title"],
               "excerpt": x["snippet"], "source": x["chapter"]} for x in res["knowledge"]]
    items += [{"kind": "material", "id": x.get("materialId") or x["id"],
               "title": x["title"], "excerpt": x["snippet"], "source": x["source"],
               "locked": not x.get("contentAvailable", True), "level": x["level"]}
              for x in res["materials"]]
    return items


def _moa_tool_search(q, user):
    """资料管理员工具：平台案例、知识和案例附件统一检索。"""
    res = search_corpus(q, limit=6, user=user)
    items = _moa_search_items(res)
    text = "\n".join(_moa_search_lines(res)) or "未检索到相关内容"
    return text[:4000], "命中 %d 条" % len(items), items


def _moa_tool_fetch(url):
    """资料管理员工具：复用 /api/fetch-url 的网页抓取。"""
    if not url:
        return "缺少 url", "缺少 url", []
    r = fetch_url({"url": url})
    if not r.get("ok"):
        return "抓取失败：" + (r.get("error") or ""), "抓取失败", []
    text = "网页《%s》（%s）内容摘录：\n%s" % (
        r.get("title") or "", r.get("finalUrl") or url, (r.get("text") or "")[:3000])
    item = {"kind": "web", "url": r.get("finalUrl") or url,
            "title": r.get("title") or url, "excerpt": (r.get("text") or "")[:500],
            "source": urllib.parse.urlsplit(r.get("finalUrl") or url).netloc}
    return text, "抓取成功：%s" % (r.get("title") or url)[:40], [item]


def _moa_tool_web_search(q):
    result = web_search({"query": q, "max_results": 6})
    if not result.get("ok"):
        return "联网检索失败：" + (result.get("error") or ""), "联网检索失败", []
    items = [{"kind": "web", "url": x["url"], "title": x["title"],
              "excerpt": x.get("content") or "",
              "source": urllib.parse.urlsplit(x["url"]).netloc} for x in result["results"]]
    lines = ["〔w%d〕%s（%s）：%s" % (i, x["title"], x["url"], x["excerpt"][:500])
             for i, x in enumerate(items, 1)]
    return "\n".join(lines) or "未检索到互联网资源", "互联网命中 %d 条" % len(items), items


def _moa_run_tool(action, args, user):
    if action == "search_corpus":
        return _moa_tool_search((args.get("q") or "").strip(), user)
    if action == "web_search":
        return _moa_tool_web_search((args.get("q") or args.get("query") or "").strip())
    return _moa_tool_fetch((args.get("url") or "").strip())


def _moa_history(payload):
    """对话历史：沿用 chat 的 [{role,content}]，服务端截断（最近 6 条、单条 1500 字）。"""
    msgs = []
    for h in (payload.get("history") or [])[-6:]:
        if not isinstance(h, dict) or h.get("role") not in ("user", "assistant"):
            continue
        content = (h.get("content") or "")[:1500]
        if content:
            msgs.append({"role": h["role"], "content": content})
    return msgs


def _moa_user_text(payload):
    """把 text + caseContext + selection 拼成给模型的用户消息。"""
    parts = []
    ctx = payload.get("caseContext") or {}
    if ctx.get("title"):
        parts.append("案例标题：" + str(ctx["title"]))
    if ctx.get("sectionTitle"):
        parts.append("当前章节：" + str(ctx["sectionTitle"]))
    if ctx.get("sectionText"):
        parts.append("章节全文：\n" + str(ctx["sectionText"])[:3000])
    elif ctx.get("bodyExcerpt"):
        parts.append("案例内容摘要：\n" + str(ctx["bodyExcerpt"])[:1500])
    citations = ctx.get("citations") or []
    if citations:
        lines = ["已关联的引用素材（含证据片段，编号即 chunk 编号）："]
        for i, c in enumerate(citations[:20], 1):
            if isinstance(c, dict):
                line = "〔%d〕%s — %s" % (c.get("n") or i, c.get("title", ""), c.get("source", ""))
                if c.get("snippet"):
                    line += "：证据「%s」" % str(c["snippet"])[:120]
                lines.append(line)
            else:
                lines.append("〔%d〕%s" % (i, c))
        parts.append("\n".join(lines))
    if payload.get("selection"):
        parts.append("用户选中的文段：\n" + str(payload["selection"])[:1500])
    parts.append("用户请求：" + (payload.get("text") or ""))
    return "\n\n".join(p for p in parts if p)


def _moa_classify(user_text):
    """无 intentHint 时由编排模型轻量分类；解析失败回落 writer。"""
    content, err = llm_call(MOA_MODEL_ORCHESTRATOR, [
        {"role": "system", "content": ORCHESTRATOR_PROMPT},
        {"role": "user", "content": user_text[:2000]},
    ], temperature=0.1)
    if content:
        obj = _extract_json(content)
        if isinstance(obj, dict) and obj.get("route") in ("librarian", "writer", "reviewer"):
            return obj["route"], str(obj.get("reason") or "")
    sys.stderr.write("[moa] 分类失败（%s），回落 writer\n" % (err or "JSON 解析失败"))
    return "writer", "分类失败，按写作处理"


def _moa_evidence_chunks(payload, user, limit=6):
    """写作/审校前的资料检索（资料管理员同套检索）：统一编号的 chunk 列表，
    = 案例已关联引用（保持原编号）+ 本次新检索命中（跳过已关联目标，续号）。
    chunk: {n, kind(knowledge|material), materialId, sec, title, source, publishedAt, grade, snippet}"""
    ctx = payload.get("caseContext") or {}
    chunks, cited = [], set()
    for i, c in enumerate((ctx.get("citations") or [])[:20], 1):
        if not isinstance(c, dict) or not c.get("target"):
            continue
        target = c["target"]
        cited.add(target)
        chunks.append({
            "n": c.get("n") or i, "kind": "material" if not str(target).startswith("kn-") else "knowledge",
            "materialId": target, "sec": c.get("sec") or "",
            "title": c.get("title") or target, "source": c.get("source") or "",
            "publishedAt": "", "grade": "",
            "snippet": c.get("snippet") or "",
        })
    q = " ".join(x for x in [(payload.get("text") or "")[:200],
                             ctx.get("title") or "", ctx.get("sectionTitle") or ""] if x).strip()
    res = search_corpus(q or "课程思政", limit=limit, user=user)
    for it in res["knowledge"]:
        if it["id"] in cited:
            continue
        chunks.append({"n": len(chunks) + 1, "kind": "knowledge", "materialId": it["id"],
                       "sec": it.get("sec", ""),
                       "title": " ".join(x for x in [it.get("chapter"), it.get("title")] if x),
                       "source": "《自然辩证法概论（2025版）》", "publishedAt": "2025",
                       "grade": "", "snippet": it.get("snippet") or ""})
    for it in res["materials"]:
        mid = it.get("materialId") or it["id"]
        if mid in cited:
            continue
        chunks.append({"n": len(chunks) + 1, "kind": "material", "materialId": mid,
                       "sec": it.get("sec", ""), "title": it.get("title") or mid,
                       "source": it.get("source") or "",
                       "publishedAt": it.get("publishedAt") or "",
                       "grade": it.get("grade") or "", "snippet": it.get("snippet") or ""})
    return chunks


def _moa_chunks_brief(chunks):
    """给模型的 chunk 编号表（写作手引用与审校员蕴含判断共用）。"""
    lines = []
    for ch in chunks:
        meta = "｜".join(x for x in [ch.get("source"), ch.get("publishedAt"),
                                    ("信源等级%s" % ch["grade"]) if ch.get("grade") else ""] if x)
        lines.append("〔%d〕%s｜%s%s：%s" % (
            ch["n"], "知识" if ch["kind"] == "knowledge" else "素材",
            ch.get("title") or ch["materialId"],
            ("（%s）" % meta) if meta else "",
            (ch.get("snippet") or "")[:160]))
    return "\n".join(lines)


def _moa_check_citations(text, chunks):
    """后处理校验：写作手输出中的〔n〕与本次实际检索/已关联的 chunk 集比对，
    对不上的降级为〔n·待核实〕并产出 risk 记录（前端落成 risk 批注），不静默放行。"""
    valid = len(chunks)
    risks = []

    def repl(m):
        n = int(m.group(1))
        if 1 <= n <= valid:
            return m.group(0)
        start = max(text.rfind("。", 0, m.start()), text.rfind("\n", 0, m.start()),
                    text.rfind("！", 0, m.start()), text.rfind("？", 0, m.start())) + 1
        end = len(text)
        for p in ("。", "！", "？", "\n"):
            q = text.find(p, m.end())
            if 0 <= q < end:
                end = q + 1
        risks.append({"n": n, "quote": text[start:end].strip()[:120],
                      "note": "引用编号〔%d〕无对应检索资料" % n})
        return "〔%d·待核实〕" % n

    out = re.sub(r"〔(\d+)〕", repl, text)
    return out, risks


def _fewshot_block():
    """反例库注入 prompt 的 few-shot 段：每类风险取 1 条（控制 token）。"""
    lines, seen = [], set()
    for e in REVIEW_COUNTEREXAMPLES:
        cat = e.get("category") or ""
        if not cat or cat in seen:
            continue
        seen.add(cat)
        lines.append("【%s】错误示例：「%s」—— %s" % (cat, e.get("bad", ""), e.get("why", "")))
    return ("\n以下是十类典型风险的错误示例，审校时对照识别同类问题：\n" + "\n".join(lines)
            if lines else "")


def _reviewer_prompt():
    return REVIEWER_PROMPT + _fewshot_block()


def _admin_names():
    """管理员姓名集合：OO 原生批注镜像时据作者名判定 kind=admin。"""
    return {u.get("name") or uid for uid, u in load_users().items() if u.get("admin")}


# --------------------------------------------------------------- docx 导出
def export_docx(payload):
    """导出成套教学材料。payload:
    title, meta{author,audience,caseType,course,mode,statusNote,footerNote},
    parts[{heading,markdown}], refs[{title,source}]
    """
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt
    import io

    title = payload.get("title") or "未命名案例"
    meta = payload.get("meta") or {}
    parts = payload.get("parts") or []
    refs = payload.get("refs") or []

    def add_markdown(doc, markdown):
        for line in markdown.splitlines():
            raw = line.rstrip()
            if not raw.strip():
                continue
            m = re.match(r"^(#{1,6})\s+(.*)$", raw)
            if m:
                doc.add_heading(re.sub(r"\*\*", "", m.group(2)), level=min(len(m.group(1)) + 1, 4))
                continue
            if re.match(r"^\s*[-*]\s+", raw):
                doc.add_paragraph(re.sub(r"^\s*[-*]\s+", "", raw), style="List Bullet")
                continue
            if re.match(r"^\s*\d+\.\s+", raw):
                doc.add_paragraph(re.sub(r"^\s*\d+\.\s+", "", raw), style="List Number")
                continue
            if raw.startswith(">"):
                p = doc.add_paragraph(raw.lstrip("> ").strip())
                p.paragraph_format.left_indent = Pt(18)
                continue
            text = re.sub(r"\[cite:[^\]]+\]", "", raw)
            text = re.sub(r"\*\*([^*]+)\*\*", r"\1", text)
            text = re.sub(r"`([^`]+)`", r"\1", text)
            doc.add_paragraph(text)

    doc = Document()
    doc.add_heading(title, level=0)
    info = doc.add_paragraph()
    info.add_run(
        "作者：%s    教学对象：%s    案例类型：%s    课程：%s    版本：%s"
        % (meta.get("author", "-"), meta.get("audience", "-"), meta.get("caseType", "-"),
           meta.get("course", "-"), meta.get("mode", "-"))
    ).font.size = Pt(9)

    for part in parts:
        heading = (part.get("heading") or "").strip()
        if heading:
            doc.add_heading(heading, level=1)
        add_markdown(doc, part.get("markdown") or "")

    if refs:
        doc.add_heading("引用与素材清单", level=1)
        for i, r in enumerate(refs, 1):
            doc.add_paragraph("[%d] %s — %s" % (i, r.get("title", ""), r.get("source", "")),
                              style="List Number")

    # 页脚追踪元数据：状态、生成时间、使用范围提示 + AI 生成标识（来源/模型/审核人）
    footer_note = meta.get("footerNote") or (
        "上海大学思政教学案例智能平台 · 生成时间：%s · 状态：%s"
        % (time.strftime("%Y-%m-%d %H:%M"), meta.get("statusNote", "未定"))
    )
    ORIGIN_NAMES = {"ai": "AI 生成", "ai_assisted": "AI 辅助", "human": "人工撰写"}
    trace = []
    if meta.get("origin"):
        trace.append("来源：" + ORIGIN_NAMES.get(meta["origin"], str(meta["origin"])))
    models = meta.get("modelVersions") or []
    if models:
        trace.append("模型：" + "/".join(str(m) for m in models))
    if meta.get("reviewedBy"):
        trace.append("审核人：" + str(meta["reviewedBy"]))
    if trace:
        footer_note += " · " + " · ".join(trace)
    for section in doc.sections:
        p = section.footer.paragraphs[0]
        p.text = footer_note
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in p.runs:
            run.font.size = Pt(8)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue(), title


# ------------------------------------------------------------ OnlyOffice 集成（WP6）
# DS（Document Server）docker 部署，JWT HS256 用标准库 hmac 实现（不引依赖）。
# docx 为编辑真源：files/cases/{caseId}.docx；blocks 保留作详情页渲染/检索/AI 上下文，
# 每次 DS 保存回调 / AI 修订插入 / 版本回滚后由 docx 解析回写，并 bump docxVer（document.key 随内容变化）。
OO_DS_URL = ENV.get("ONLYOFFICE_DS_URL", "").rstrip("/")       # 浏览器可达的 DS 地址（api.js 来源）
OO_SERVER_URL = ENV.get("ONLYOFFICE_SERVER_URL", "").rstrip("/")  # DS 容器可达的本服务地址
OO_JWT_SECRET = ENV.get("ONLYOFFICE_JWT_SECRET", "")
CASES_DOCX_DIR = os.path.join(FILES_DIR, "cases")
OO_LOCK = threading.Lock()  # docx 文件读写串行化
OO_AUTHOR = "AI 助手"       # AI 修订（w:ins/w:del）的作者署名


def oo_configured():
    return bool(OO_DS_URL and OO_JWT_SECRET)


def _oo_jwt_sign(payload_obj):
    """JWT HS256（OnlyOffice 官方签名格式）：header/payload/signature 三段 b64url。"""
    header = _b64url(b'{"alg":"HS256","typ":"JWT"}')
    payload = _b64url(json.dumps(payload_obj, separators=(",", ":"),
                                 ensure_ascii=False).encode("utf-8"))
    sig = _b64url(hmac.new(OO_JWT_SECRET.encode("utf-8"),
                           (header + "." + payload).encode("ascii"),
                           hashlib.sha256).digest())
    return header + "." + payload + "." + sig


def _oo_jwt_verify(token):
    """校验 DS 回传 JWT 签名，返回 payload；无效返回 None。"""
    if not OO_JWT_SECRET or not token or token.count(".") != 2:
        return None
    header, payload, sig = token.split(".")
    want = _b64url(hmac.new(OO_JWT_SECRET.encode("utf-8"),
                            (header + "." + payload).encode("ascii"),
                            hashlib.sha256).digest())
    if not hmac.compare_digest(sig, want):
        return None
    try:
        return json.loads(_b64url_decode(payload))
    except Exception:
        return None


def case_docx_path(cid):
    return os.path.join(CASES_DOCX_DIR, cid + ".docx")


def blocks_to_docx(blocks, path):
    """blocks → docx：h2→Heading 1，p→正文，ul/ol→List Bullet/Number（一项一段），
    quote→Quote 样式；〔n〕上标引用标记保留为字面文本。"""
    from docx import Document
    os.makedirs(os.path.dirname(path), exist_ok=True)
    doc = Document()
    for b in blocks or []:
        kind = b.get("kind")
        lines = [ln.strip() for ln in str(b.get("text") or "").split("\n") if ln.strip()]
        if not lines:
            continue
        if kind == "h2":
            doc.add_heading(lines[0], level=1)
        elif kind == "ul":
            for ln in lines:
                doc.add_paragraph(ln, style="List Bullet")
        elif kind == "ol":
            for ln in lines:
                doc.add_paragraph(ln, style="List Number")
        elif kind == "quote":
            for ln in lines:
                try:
                    doc.add_paragraph(ln, style="Quote")
                except KeyError:
                    doc.add_paragraph(ln)
        else:
            for ln in lines:
                doc.add_paragraph(ln)
    doc.save(path)


def _para_full_text(p):
    """段落全文：含 w:ins 修订内文本（python-docx p.text 会漏掉），排除 w:del 已删文本。"""
    from docx.oxml.ns import qn
    out = []
    for node in p._p.iter():
        if node.tag == qn("w:t"):
            out.append(node.text or "")
    return "".join(out)


def docx_to_blocks(path):
    """docx → blocks：Heading*→h2，List Bullet/Number 连续段合并为一个 ul/ol 块（一项一行），
    Quote 连续段合并为 quote 块（一行一段），其余非空段→p。"""
    from docx import Document
    doc = Document(path)
    blocks, acc_kind, acc = [], None, []

    def flush():
        nonlocal acc_kind, acc
        if acc_kind and acc:
            blocks.append({"kind": acc_kind, "text": "\n".join(acc)})
        acc_kind, acc = None, []

    for p in doc.paragraphs:
        style = p.style.name if p.style is not None else ""
        text = _para_full_text(p).strip()
        if style.startswith("Heading"):
            flush()
            if text:
                blocks.append({"kind": "h2", "text": text})
            continue
        kind = ("ul" if "Bullet" in style else
                "ol" if "Number" in style else
                "quote" if "Quote" in style else None)
        if kind:
            if acc_kind != kind:
                flush()
                acc_kind = kind
            if text:
                acc.append(text)
            continue
        flush()
        if text:
            blocks.append({"kind": "p", "text": text})
    flush()
    return blocks


def ensure_case_docx(cid, blocks):
    """无 docx 时从 blocks 生成（老案例迁移/文件丢失兜底）；返回是否新生成。"""
    path = case_docx_path(cid)
    if os.path.isfile(path):
        return False
    blocks_to_docx(blocks, path)
    return True


# ------------------------------------------------------------ OO 原生批注（comments.xml）
# 批注统一载体（OO 为唯一真源）：侧栏/AI 审校批注注入 docx comment，
# DS 保存回调时解析 comments.xml 镜像回 annotations 表（ooId 关联，docx 中删失即 resolved）。
# OO 9.x 保存后会多写一套私有批注流（schemas.onlyoffice.com/commentsDocument 等 parts），
# 再次加载时与标准流按 w:id 去重合并——注入必须双写且 id 跨流唯一，否则注入的批注会被吃掉。
OO_PRIVATE_COMMENTS = "word/commentsDocument.xml"


def _zip_read_member(path, member):
    import zipfile
    try:
        with zipfile.ZipFile(path) as z:
            return z.read(member)
    except Exception:
        return None


def _zip_replace_member(path, member, data):
    """重写 zip 中指定成员（OO 私有批注流 python-docx 不管，需 zip 级手术）。"""
    import zipfile
    tmp = path + ".tmp"
    with zipfile.ZipFile(path) as zin, \
            zipfile.ZipFile(tmp, "w", zipfile.ZIP_DEFLATED) as zout:
        for item in zin.infolist():
            zout.writestr(item, data if item.filename == member
                          else zin.read(item.filename))
    os.replace(tmp, path)


def _oo_private_comment_ids(path):
    """OO 私有批注流中的 w:id 集合（无该 part 返回空集）。"""
    raw = _zip_read_member(path, OO_PRIVATE_COMMENTS)
    if not raw:
        return set()
    from docx.oxml import parse_xml
    from docx.oxml.ns import qn
    try:
        root = parse_xml(raw)
        return {int(el.get(qn("w:id"))) for el in root.iter(qn("w:comment"))
                if el.get(qn("w:id")) is not None}
    except Exception:
        return set()


def _oo_private_add_comment(path, comment_el, new_id):
    """把新批注元素复制进 OO 私有批注流（id 与标准流一致；无该 part 跳过）。"""
    import copy
    from docx.oxml.ns import qn
    raw = _zip_read_member(path, OO_PRIVATE_COMMENTS)
    if not raw:
        return False
    from docx.oxml import parse_xml
    try:
        root = parse_xml(raw)
        el = copy.deepcopy(comment_el)
        el.set(qn("w:id"), str(new_id))
        root.append(el)
        from lxml import etree
        _zip_replace_member(path, OO_PRIVATE_COMMENTS, etree.tostring(
            root, xml_declaration=True, encoding="UTF-8", standalone=True))
        return True
    except Exception as e:
        sys.stderr.write("[oo] 私有批注流写入失败：%s\n" % e)
        return False


def _oo_private_remove_comment(path, oo_id):
    """从 OO 私有批注流移除指定 id 的批注；返回是否有该 part。"""
    from docx.oxml.ns import qn
    raw = _zip_read_member(path, OO_PRIVATE_COMMENTS)
    if not raw:
        return False
    from docx.oxml import parse_xml
    try:
        root = parse_xml(raw)
        hit = False
        for el in list(root.iter(qn("w:comment"))):
            if el.get(qn("w:id")) == str(oo_id):
                el.getparent().remove(el)
                hit = True
        if hit:
            from lxml import etree
            _zip_replace_member(path, OO_PRIVATE_COMMENTS, etree.tostring(
                root, xml_declaration=True, encoding="UTF-8", standalone=True))
        return True
    except Exception:
        return True

def _oo_run_split(run, offset):
    """在 run 文本 offset 处拆成两个 run（深拷贝保留 rPr），返回后半 run。"""
    import copy
    from docx.text.run import Run
    text = run.text
    if offset <= 0 or offset >= len(text):
        return None
    new_r = copy.deepcopy(run._r)
    run.text = text[:offset]
    run._r.addnext(new_r)
    nr = Run(new_r, run._parent)
    nr.text = text[offset:]
    return nr


def _oo_split_at(p, offset):
    """让段落 p 的拼接文本 offset 处落在 run 边界（必要时拆 run）；
    返回从 offset 开始的 run；offset 超出段末返回 None。"""
    pos = 0
    for r in p.runs:
        t = r.text
        if offset == pos:
            return r
        if pos < offset < pos + len(t):
            return _oo_run_split(r, offset - pos)
        pos += len(t)
    return None


def _oo_span(p_start, off_start, p_end, off_end):
    """构造锚定区间（按段拼接文本偏移）：先拆终点再拆起点（同段时保持偏移有效），
    返回 (first_run, last_run)；退化情况锚到 start_run 本身。"""
    from docx.text.run import Run
    from docx.oxml.ns import qn
    end_next = _oo_split_at(p_end, off_end)
    start_run = _oo_split_at(p_start, off_start) or (p_start.runs[0] if p_start.runs else None)
    if start_run is None:
        return None
    if end_next is not None:
        prev = end_next._r.getprevious()
        while prev is not None and prev.tag != qn("w:r"):
            prev = prev.getprevious()
        if prev is not None:
            return (start_run, Run(prev, end_next._parent))
        return (start_run, start_run)
    runs = p_end.runs
    return (start_run, runs[-1] if runs else start_run)


def _oo_anchor_runs(doc, quote):
    """按 quote 定位锚定 run 区间（先段内精确匹配，再全文空白归一化匹配，支持跨段）；
    找不到返回 None。"""
    quote = (quote or "").strip()
    if not quote:
        return None
    paras = [p for p in doc.paragraphs if p.runs]
    texts = ["".join(r.text for r in p.runs) for p in paras]
    for i, full in enumerate(texts):
        idx = full.find(quote)
        if idx >= 0:
            return _oo_span(paras[i], idx, paras[i], idx + len(quote))
    norm, pmap = "", []  # 归一化串 + 下标→(段, 段内字符) 映射
    for i, full in enumerate(texts):
        for j, ch in enumerate(full):
            if not ch.isspace():
                norm += ch
                pmap.append((i, j))
    nq = "".join(ch for ch in quote if not ch.isspace())
    if not nq:
        return None
    k = norm.find(nq)
    if k < 0:
        return None
    (i1, j1), (i2, j2) = pmap[k], pmap[k + len(nq) - 1]
    return _oo_span(paras[i1], j1, paras[i2], j2 + 1)


def oo_inject_comment(cid, quote, text, author):
    """把一条批注作为 OO 原生批注注入 docx：锚定 quote 文本区间（找不到锚到末段）。
    id 跨标准/私有批注流唯一，并双写私有流（OO 9.x 合并去重语义，否则注入批注加载时被吃）。
    返回 (comment_id, 是否精确锚定, error)。"""
    from docx import Document
    from docx.oxml.ns import qn
    path = case_docx_path(cid)
    if not os.path.isfile(path):
        return None, False, "案例 docx 不存在"
    with OO_LOCK:
        try:
            doc = Document(path)
            span = _oo_anchor_runs(doc, quote or "")
            anchored = bool(span)
            if not span:
                paras = [p for p in doc.paragraphs if p.runs]
                if not paras:
                    return None, False, "docx 为空，无法锚定批注"
                span = (paras[-1].runs[0], paras[-1].runs[-1])
            comment = doc.add_comment(span, text=text, author=author or "批注", initials=None)
            old_id = comment.comment_id
            new_id = max([old_id] + list(_oo_private_comment_ids(path))) + 1 \
                if old_id in _oo_private_comment_ids(path) else old_id
            if new_id != old_id:
                comment._comment_elm.set(qn("w:id"), str(new_id))
                for el in doc.element.body.iter():
                    if el.tag in (qn("w:commentRangeStart"), qn("w:commentRangeEnd"),
                                  qn("w:commentReference")) \
                            and el.get(qn("w:id")) == str(old_id):
                        el.set(qn("w:id"), str(new_id))
            doc.save(path)
            _oo_private_add_comment(path, comment._comment_elm, new_id)
            return new_id, anchored, None
        except Exception as e:
            return None, False, "批注注入 docx 失败：%s" % e


def oo_remove_comment(cid, oo_id):
    """从 docx 移除 id=oo_id 的批注（range 起止标记 + reference run + comments.xml 条目）。
    返回是否找到并移除。"""
    from docx import Document
    from docx.oxml.ns import qn
    path = case_docx_path(cid)
    if not os.path.isfile(path):
        return False
    with OO_LOCK:
        try:
            doc = Document(path)
            sid = str(oo_id)
            found = False
            for el in list(doc.element.body.iter()):
                if el.tag in (qn("w:commentRangeStart"), qn("w:commentRangeEnd")) \
                        and el.get(qn("w:id")) == sid:
                    el.getparent().remove(el)
                    found = True
                elif el.tag == qn("w:commentReference") and el.get(qn("w:id")) == sid:
                    r = el.getparent()
                    if r is not None and r.tag == qn("w:r"):
                        r.getparent().remove(r)
                    found = True
            try:
                cel = doc.comments._comments_elm.get_comment_by_id(int(oo_id))
            except Exception:
                cel = None
            if cel is not None:
                cel.getparent().remove(cel)
                found = True
            if found:
                doc.save(path)
            return found
        except Exception as e:
            sys.stderr.write("[oo] 批注移除失败 %s #%s: %s\n" % (cid, oo_id, e))
            return False


def oo_parse_comments(path):
    """解析 docx 全部 OO 批注：comments.xml（首段=正文，其余段=回复）
    + document.xml range 提取 quote。返回 [{ooId, author, date, text, quote, replies}]。"""
    from docx import Document
    from docx.oxml.ns import qn
    if not os.path.isfile(path):
        return []
    try:
        doc = Document(path)
    except Exception:
        return []
    quotes, active = {}, set()
    for el in doc.element.body.iter():
        if el.tag == qn("w:commentRangeStart"):
            active.add(el.get(qn("w:id")))
        elif el.tag == qn("w:commentRangeEnd"):
            active.discard(el.get(qn("w:id")))
        elif el.tag == qn("w:t") and active:
            for k in active:
                quotes.setdefault(k, []).append(el.text or "")
    quotes = {k: "".join(v) for k, v in quotes.items()}
    out = []
    try:
        comments = list(doc.comments)
    except Exception:
        return out
    for cm in comments:
        oid = str(cm.comment_id)
        paras = [p.text for p in cm.paragraphs]
        text = paras[0] if paras else ""
        replies = [{"byName": cm.author or "批注", "text": t, "at": ""}
                   for t in paras[1:] if t.strip()]
        ts = cm.timestamp
        out.append({"ooId": oid, "author": cm.author or "",
                    "date": ts.strftime("%Y-%m-%d %H:%M") if ts else "",
                    "text": text, "quote": quotes.get(oid, "")[:200],
                    "replies": replies})
    return out


def migrate_cases_docx():
    """启动迁移：为全部存量案例补齐 docx（种子 4 篇含在内）。"""
    n = 0
    with OO_LOCK:
        for c in CASEDB.list_cases({"id": "system", "admin": True}):
            if ensure_case_docx(c["id"], c.get("blocks")):
                n += 1
    return n


def _oo_rev_mark(p_el, tag, author, date, rev_id):
    """把段落内的 run 包进 w:ins/w:del 修订标记；w:del 时 w:t 改 w:delText。"""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    el = OxmlElement(tag)
    el.set(qn("w:id"), str(rev_id))
    el.set(qn("w:author"), author)
    el.set(qn("w:date"), date)
    for r in p_el.findall(qn("w:r")):
        if tag == "w:del":
            for t in r.findall(qn("w:t")):
                t.tag = qn("w:delText")
        p_el.remove(r)
        el.append(r)
    p_el.append(el)


_OO_REV_ID = [1000]


def _oo_rev_next():
    _OO_REV_ID[0] += 1
    return _OO_REV_ID[0]


def oo_insert_revision(cid, mode, text, sec_title=""):
    """AI 内容以 track-changes 修订形式写入 docx（教师在 OnlyOffice 里接受/拒绝）：
    append=插到 secTitle 节末（无 secTitle 插文末）；newsec=文末新 h2 节；
    replace=secTitle 节正文整体 w:del 标记删除 + 新内容 w:ins 紧随其后。
    全部新段落以 w:ins 包裹。返回 (blocks, error)；成功后由调用方回写 blocks 并 bump docxVer。"""
    from docx import Document
    from docx.oxml.ns import qn
    path = case_docx_path(cid)
    if not os.path.isfile(path):
        return None, "案例 docx 不存在"
    paras = [ln.strip() for ln in str(text or "").split("\n") if ln.strip()]
    if mode == "newsec" and paras and paras[0].replace("#", "").strip() == (sec_title or "").strip():
        pass  # 模型复述节标题时保留为首行标题
    if not paras:
        return None, "没有可写入的内容"
    date = time.strftime("%Y-%m-%dT%H:%M:%SZ", time.gmtime())
    with OO_LOCK:
        doc = Document(path)
        body_pars = list(doc.paragraphs)
        # 定位节：h2 文本与 sec_title 相同的段落；节末=下一个 Heading 或文末
        sec_head, sec_end = None, None
        if sec_title and mode in ("append", "replace"):
            for i, p in enumerate(body_pars):
                style = p.style.name if p.style is not None else ""
                if style.startswith("Heading") and _para_full_text(p).strip() == sec_title:
                    sec_head = i
                    for j in range(i + 1, len(body_pars)):
                        st = body_pars[j].style.name if body_pars[j].style is not None else ""
                        if st.startswith("Heading"):
                            sec_end = j
                            break
                    break
            if sec_head is None:
                return None, "docx 中未找到小节「%s」" % sec_title

        def new_par(txt, style=None, tracked=True):
            np = doc.add_paragraph(txt)  # 先挂到文末，再移动到目标位置
            if style:
                try:
                    np.style = doc.styles[style]
                except KeyError:
                    pass
            if tracked:
                _oo_rev_mark(np._p, "w:ins", OO_AUTHOR, date, _oo_rev_next())
            return np

        anchor = None  # 插入点：插到 anchor 段之前；None=文末
        if mode == "newsec":
            head = paras[0].lstrip("#").strip()[:60]
            rest = paras[1:] if len(paras) > 1 else paras
            new_par(head, "Heading 1")
            for t in rest:
                new_par(t)
        elif mode == "replace" and sec_head is not None:
            # 节正文（标题之后、下一标题之前）整体标记删除，新内容插在下一标题前
            end = sec_end if sec_end is not None else len(body_pars)
            for p in body_pars[sec_head + 1:end]:
                if _para_full_text(p).strip():
                    _oo_rev_mark(p._p, "w:del", OO_AUTHOR, date, _oo_rev_next())
            anchor = body_pars[sec_end]._p if sec_end is not None else None
            news = [new_par(t) for t in paras]
            if anchor is not None:
                for np in news:
                    anchor.addprevious(np._p)
        else:  # append（有 secTitle 插节末，否则插文末）
            if sec_head is not None:
                anchor = body_pars[sec_end]._p if sec_end is not None else None
            news = [new_par(t) for t in paras]
            if anchor is not None:
                for np in news:
                    anchor.addprevious(np._p)
        doc.save(path)
        return docx_to_blocks(path), None


def oo_download(url):
    """从 DS 下载文件（callback 的 body.url）；JWT 头优先，失败降级裸 GET。"""
    for headers in ({"Authorization": "Bearer " + _oo_jwt_sign({})},
                    {"User-Agent": UA}):
        try:
            req = urllib.request.Request(url, headers=headers)
            with urllib.request.urlopen(req, timeout=30) as resp:
                return resp.read(), None
        except Exception:
            continue
    return None, "从 Document Server 下载保存文件失败"


# ------------------------------------------------------------ OnlyOffice API
def _oo_doc_ver(cid):
    data = CASEDB.get_case_data(cid) or {}
    return int(data.get("docxVer") or 1)


def api_oo_config(user, cid):
    """DocEditor 初始化 config：document.url 指回本服务文件端点，callbackUrl 指回 callback，
    key=caseId-docxVer（内容变化即变），permissions 按 owner/admin/只读区分；整体 JWT 签名。"""
    err = _need_login(user)
    if err:
        return err
    if not oo_configured():
        return {"ok": False, "error": "服务端未配置 ONLYOFFICE_*"}, 503
    c = CASEDB.get_case(cid, user)
    if not c:
        return {"ok": False, "error": "案例不存在或无权查看"}, 404
    base = OO_SERVER_URL or ""
    if not base:
        return {"ok": False, "error": "服务端未配置 ONLYOFFICE_SERVER_URL"}, 503
    with OO_LOCK:
        ensure_case_docx(cid, c.get("blocks"))
    can_edit = c["ownerId"] == user["id"] and c.get("status") == "draft"
    can_comment = bool(user.get("admin"))
    ver = _oo_doc_ver(cid)
    config = {
        "document": {
            "fileType": "docx",
            "key": "%s-v%d" % (cid, ver),
            "title": (c.get("title") or cid) + ".docx",
            "url": "%s/api/onlyoffice/file/%s" % (base, urllib.parse.quote(cid)),
            "permissions": {
                "edit": can_edit,
                "comment": can_edit or can_comment,
                "review": can_edit,
                "download": True,
                "print": True,
            },
        },
        "documentType": "word",
        "editorConfig": {
            "callbackUrl": "%s/api/onlyoffice/callback?caseId=%s" % (base, urllib.parse.quote(cid)),
            "lang": "zh-CN",
            "mode": "edit" if can_edit else "view",
            "user": {"id": user["id"], "name": user.get("name") or user["id"]},
            "customization": {
                "autosave": True,
                "forcesave": True,
                "comments": can_edit or can_comment,
                "compactToolbar": False,
            },
        },
    }
    config["token"] = _oo_jwt_sign(config)
    return {"ok": True, "config": config, "docxVer": ver,
            "apiJs": OO_DS_URL + "/web-apps/apps/api/documents/api.js",
            "editable": can_edit}, 200


def api_oo_file(handler, cid):
    """DS 拉取 docx：仅接受带有效 JWT 签名的请求（DS 下载时自动附带）。"""
    if not oo_configured():
        return None, {"ok": False, "error": "服务端未配置 ONLYOFFICE_*"}, 404
    ah = handler.headers.get("Authorization") or ""
    tok = ah[7:].strip() if ah.startswith("Bearer ") else ""
    if _oo_jwt_verify(tok) is None:
        return None, {"ok": False, "error": "JWT 校验失败"}, 403
    path = case_docx_path(cid)
    if not os.path.isfile(path):
        return None, {"ok": False, "error": "案例 docx 不存在"}, 404
    with OO_LOCK, open(path, "rb") as f:
        data = f.read()
    return data, None, 200


def api_oo_callback(handler, query, payload):
    """DS 保存回调：status 2/6=保存（下载 body.url 落盘 + 回写 blocks + bump key）；
    status 4=关闭无修改。请求须带有效 JWT（payload 为回调 body）。恒回 {"error":0} 防 DS 重试。"""
    cid = (query.get("caseId") or [""])[0].strip()
    if not oo_configured():
        return {"error": 0}, 200
    ah = handler.headers.get("Authorization") or ""
    tok = ah[7:].strip() if ah.startswith("Bearer ") else ""
    if _oo_jwt_verify(tok) is None:
        sys.stderr.write("[oo] callback JWT 校验失败 %s\n" % cid)
        return {"error": 0}, 200
    status = payload.get("status")
    if status in (2, 6) and payload.get("url") and cid:
        raw, derr = oo_download(payload["url"])
        if derr:
            sys.stderr.write("[oo] callback 下载失败 %s: %s\n" % (cid, derr))
            return {"error": 0}, 200
        with OO_LOCK:
            path = case_docx_path(cid)
            os.makedirs(os.path.dirname(path), exist_ok=True)
            with open(path, "wb") as f:
                f.write(raw)
            try:
                blocks = docx_to_blocks(path)
            except Exception as e:
                sys.stderr.write("[oo] callback 解析 docx 失败 %s: %s\n" % (cid, e))
                return {"error": 0}, 200
            c, e = CASEDB.apply_blocks_from_docx(cid, blocks)
        if e:
            sys.stderr.write("[oo] callback 回写失败 %s: %s\n" % (cid, e))
        else:
            n = CASEDB.mirror_oo_comments(cid, oo_parse_comments(path), _admin_names())
            mark_search_dirty()
            sys.stderr.write("[oo] 已保存 %s（blocks %d 块，docxVer %d，批注镜像 %d 条）\n"
                             % (cid, len(blocks), _oo_doc_ver(cid), n))
    return {"error": 0}, 200


def api_oo_insert(user, payload):
    """AI 修订插入（track-changes）：Copilot 起草/改写采纳时调用；
    成功后 docxVer 已 bump，前端用新 config 重建编辑器强制 DS 重载。"""
    err = _need_login(user)
    if err:
        return err
    if not oo_configured():
        return {"ok": False, "error": "服务端未配置 ONLYOFFICE_*"}, 503
    cid = (payload.get("caseId") or "").strip()
    c = CASEDB.get_case(cid, user)
    if not c:
        return {"ok": False, "error": "案例不存在或无权查看"}, 404
    if c["ownerId"] != user["id"] or c.get("status") != "draft":
        return {"ok": False, "error": "仅作者本人的草稿可写入 AI 修订"}, 403
    mode = (payload.get("mode") or "append").strip()
    if mode not in ("append", "newsec", "replace"):
        return {"ok": False, "error": "mode 取值无效（append/newsec/replace）"}, 400
    with OO_LOCK:
        ensure_case_docx(cid, c.get("blocks"))
    blocks, ierr = oo_insert_revision(cid, mode, payload.get("text") or "",
                                      (payload.get("secTitle") or "").strip())
    if ierr:
        return {"ok": False, "error": ierr}, 400
    c2, e = CASEDB.apply_blocks_from_docx(cid, blocks)
    if e:
        return {"ok": False, "error": e}, 404
    mark_search_dirty()
    return {"ok": True, "case": c2, "docxVer": _oo_doc_ver(cid)}, 200


def sync_docx_from_blocks(cid, blocks, only_if_exists=False):
    """blocks 直写路径（兼容编辑器 PATCH / 版本回滚）后同步 docx 并 bump key。"""
    if not oo_configured():
        return
    if only_if_exists and not os.path.isfile(case_docx_path(cid)):
        return  # 从未用 OnlyOffice 打开过的案例不预生成，docx 首开时惰性迁移
    with OO_LOCK:
        try:
            blocks_to_docx(blocks, case_docx_path(cid))
        except Exception as e:
            sys.stderr.write("[oo] docx 同步失败 %s: %s\n" % (cid, e))
            return
    CASEDB.apply_blocks_from_docx(cid, blocks)


def save_version_docx(cid, vid):
    """版本快照：复制当前 docx 到 files/cases/versions/{cid}-{vid}.docx。"""
    src = case_docx_path(cid)
    if not os.path.isfile(src):
        return
    dst_dir = os.path.join(CASES_DOCX_DIR, "versions")
    os.makedirs(dst_dir, exist_ok=True)
    with OO_LOCK, open(src, "rb") as f:
        data = f.read()
    with open(os.path.join(dst_dir, "%s-%s.docx" % (cid, vid)), "wb") as f:
        f.write(data)


def delete_case_docx(cid):
    """删除案例时联动清理 docx 与版本副本。"""
    for p in [case_docx_path(cid)] + [
            os.path.join(CASES_DOCX_DIR, "versions", n)
            for n in (os.listdir(os.path.join(CASES_DOCX_DIR, "versions"))
                      if os.path.isdir(os.path.join(CASES_DOCX_DIR, "versions")) else [])
            if n.startswith(cid + "-")]:
        try:
            os.remove(p)
        except OSError:
            pass


# --------------------------------------------------------------- HTTP 处理
class Handler(BaseHTTPRequestHandler):
    protocol_version = "HTTP/1.1"

    def log_message(self, fmt, *args):
        sys.stderr.write("[app] %s - %s\n" % (self.address_string(), fmt % args))

    # 工具 ------------------------------------------------------------
    def _send_json(self, obj, status=200):
        body = json.dumps(obj, ensure_ascii=False).encode("utf-8")
        self.send_response(status)
        self.send_header("Content-Type", "application/json; charset=utf-8")
        self.send_header("Content-Length", str(len(body)))
        self.end_headers()
        self.wfile.write(body)

    def _read_json(self):
        length = int(self.headers.get("Content-Length") or 0)
        raw = self.rfile.read(length) if length else b"{}"
        try:
            return json.loads(raw.decode("utf-8"))
        except Exception:
            return {}

    # 路由 ------------------------------------------------------------
    def _sse_chunk(self, data):
        """写一个 chunked 编码块并立即 flush。"""
        self.wfile.write(b"%x\r\n" % len(data) + data + b"\r\n")
        self.wfile.flush()

    # MoA SSE 帧 -----------------------------------------------------
    def _sse_begin(self):
        self.send_response(200)
        self.send_header("Content-Type", "text/event-stream; charset=utf-8")
        self.send_header("Cache-Control", "no-cache")
        self.send_header("Transfer-Encoding", "chunked")
        self.end_headers()

    def _sse_frame(self, obj):
        """逐帧 data: {json}\n\n，与前端契约一致。"""
        self._sse_chunk(("data: " + json.dumps(obj, ensure_ascii=False) + "\n\n").encode("utf-8"))

    def _sse_end(self):
        try:
            self.wfile.write(b"0\r\n\r\n")  # chunked 结束标记
            self.wfile.flush()
        except Exception:
            pass

    def _sse_role(self, agent, model):
        meta = MOA_AGENT_META[agent]
        self._sse_frame({"type": "role", "agent": agent, "label": meta["label"],
                         "model": model, "skill": meta["skill"]})

    def _ai_agent(self, payload):
        """统一 AI 入口：主Agent 分派 → 资料管理员/写作手（双候选）/内容审校员。"""
        if not AI_BASE_URL or not AI_API_KEY:
            return self._send_json({"ok": False, "error": "服务端未配置 AI_BASE_URL / AI_API_KEY"})
        ctx = payload.get("caseContext") or {}
        if not ((payload.get("text") or "").strip()
                or payload.get("selection") or ctx.get("sectionText")):
            return self._send_json({"ok": False, "error": "缺少请求内容 text"}, 400)
        user = auth_user(self)
        prefs = CASEDB.get_prefs(user["id"]) if (user and CASEDB is not None) else {}
        self._sse_begin()
        try:
            hint = (payload.get("intentHint") or "").strip()
            route = INTENT_ROUTES.get(hint)
            user_text = _moa_user_text(payload)
            if not route:
                self._sse_role("orchestrator", MOA_MODEL_ORCHESTRATOR)
                route, _reason = _moa_classify(user_text)
            history = _moa_history(payload)
            if route == "librarian":
                self._agent_librarian(user_text, history, user)
            elif route == "reviewer":
                self._agent_reviewer(user_text, history, payload, user)
            else:
                self._agent_writer(user_text, history, payload, user, prefs)
            self._sse_frame({"type": "done"})
        except (BrokenPipeError, ConnectionResetError):
            pass
        except Exception as e:
            try:
                self._sse_frame({"type": "error", "message": "编排失败: %s" % e})
            except Exception:
                pass
        finally:
            self._sse_end()

    def _agent_librarian(self, user_text, history, user):
        """资料管理员：内部 JSON contract 工具循环（≤3 次工具调用后强制 final）。"""
        model = MOA_MODEL_LIBRARIAN
        self._sse_role("librarian", model)
        msgs = ([{"role": "system", "content": LIBRARIAN_PROMPT}] + history
                + [{"role": "user", "content": user_text}])
        final_text, last_content, last_summary, tool_calls = None, "", "", 0
        for _round in range(4):
            content, err = llm_call(model, msgs, temperature=0.3)
            if err:
                self._sse_frame({"type": "error", "message": err})
                return
            last_content = content or last_content
            obj = _extract_json(content)
            action = obj.get("action") if isinstance(obj, dict) else None
            if action in ("search_corpus", "web_search", "fetch_url") and tool_calls < 3:
                tool_calls += 1
                args = obj.get("args") or {}
                tool_text, summary, items = _moa_run_tool(action, args, user)
                self._sse_frame({"type": "tool", "agent": "librarian",
                                 "tool": action, "args": args, "summary": summary,
                                 "items": items})
                last_summary = summary
                msgs.append({"role": "assistant", "content": content})
                msgs.append({"role": "user",
                             "content": "工具结果：\n" + tool_text + "\n\n请继续（输出下一个 JSON）。"})
                continue
            if action == "final":
                final_text = str(obj.get("content") or "").strip() or (content or "")
                break
            nudge = ("工具调用已达上限，" if action in ("search_corpus", "web_search", "fetch_url")
                     else "输出无法解析，") + "请立即输出最终答复 JSON：{\"action\":\"final\",\"content\":\"...\"}"
            msgs.append({"role": "assistant", "content": content or ""})
            msgs.append({"role": "user", "content": nudge})
        if final_text is None:
            obj = _extract_json(last_content)
            final_text = str(obj.get("content") or "").strip() \
                if isinstance(obj, dict) and obj.get("action") == "final" else ""
        if not final_text:
            final_text = (last_summary + "。请从检索结果中选择需要引用的内容。") \
                if last_summary else "没有形成可用结果，请调整检索词后重试。"
        for piece in _chunk_text(final_text):
            self._sse_frame({"type": "token", "agent": "librarian",
                             "which": "main", "text": piece})
        self._sse_frame({"type": "result", "kind": "text", "text": final_text})

    def _agent_writer(self, user_text, history, payload, user, prefs=None):
        """写作手：先做资料检索（chunk 编号表注入 prompt），ThreadPoolExecutor 并行双候选，
        成稿经引用后处理校验（对不上 chunk 的〔n〕降级「待核实」+ risk 记录）。
        教师有生成偏好（WP4b）时注入偏好段；禁用词命中在候选 risks 里加警示（前端转 risk 批注）。"""
        self._sse_role("librarian", MOA_MODEL_LIBRARIAN)
        chunks = _moa_evidence_chunks(payload, user)
        self._sse_frame({"type": "tool", "agent": "librarian", "tool": "search_corpus",
                         "args": {"q": (payload.get("text") or "")[:60]},
                         "summary": "可用资料 chunk %d 条（含已关联引用）" % len(chunks)})
        if chunks:
            user_text += ("\n\n【本次可用资料 chunk（事实与理论只能依据下表，〔n〕为引用编号）】\n"
                          + _moa_chunks_brief(chunks))
        pref_block = prefs_prompt_block(prefs)
        if pref_block:
            user_text += "\n\n" + pref_block
        sys.stderr.write("[prefs] 写作手 prompt %s教师偏好段\n" % ("含" if pref_block else "不含"))
        self._sse_role("writer", AI_DEFAULT_MODEL)
        msgs = ([{"role": "system", "content": WRITER_PROMPT}] + history
                + [{"role": "user", "content": user_text}])
        q = queue.Queue()

        def job(which, model):
            full, err = llm_stream_call(
                model, msgs, on_delta=lambda p: q.put(("token", which, p)))
            q.put(("end", which, model, full, err))

        results, ended = {}, 0
        with ThreadPoolExecutor(max_workers=2) as ex:
            ex.submit(job, "main", AI_DEFAULT_MODEL)
            ex.submit(job, "alt", MOA_MODEL_WRITER_ALT)
            while ended < 2:
                item = q.get()
                if item[0] == "token":
                    self._sse_frame({"type": "token", "agent": "writer",
                                     "which": item[1], "text": item[2]})
                else:
                    _, which, model, full, err = item
                    results[which] = {
                        "model": model,
                        "text": full if full is not None else "（该候选生成失败：%s）" % err,
                    }
                    ended += 1
        for r in results.values():
            if r.get("text"):
                r["text"], r["risks"] = _moa_check_citations(r["text"], chunks)
                for w in banned_word_hits(r["text"], prefs):
                    r["risks"].append({"n": 0, "quote": w, "banned": True,
                                       "note": "生成结果命中你的禁用词「%s」，建议修改后再采纳" % w})
        fallback = {"model": "", "text": ""}
        self._sse_frame({"type": "result", "kind": "candidates",
                         "main": results.get("main") or fallback,
                         "alt": results.get("alt") or fallback,
                         "chunks": chunks})

    def _agent_reviewer(self, user_text, history, payload, user):
        """内容审校员：检索 chunk 表作为引用蕴含判断依据；按标准清单输出严格 JSON 数组。"""
        self._sse_role("reviewer", MOA_MODEL_REVIEWER)
        chunks = _moa_evidence_chunks(payload, user)
        if chunks:
            user_text += ("\n\n【引用 chunk 表（文中〔n〕对应的资料原文，做引用蕴含判断的依据）】\n"
                          + _moa_chunks_brief(chunks))
        msgs = ([{"role": "system", "content": _reviewer_prompt()}] + history
                + [{"role": "user", "content": user_text}])
        content, err = llm_call(MOA_MODEL_REVIEWER, msgs, temperature=0.2)
        if err:
            self._sse_frame({"type": "error", "message": err})
            return
        arr = _extract_json(content, want_array=True)
        items = []
        if isinstance(arr, list):
            for it in arr:
                if not isinstance(it, dict):
                    continue
                status = it.get("status")
                item = {
                    "standard": str(it.get("standard") or "未命名标准"),
                    "status": status if status in ("pass", "risk", "confirm") else "confirm",
                    "note": str(it.get("note") or ""),
                }
                if it.get("ref"):
                    item["ref"] = str(it["ref"])
                items.append(item)
        if not items:
            items = [{"standard": "审校结果解析", "status": "confirm",
                      "note": (content or "").strip() or "审校结果为空，需人工复核"}]
        self._sse_frame({"type": "result", "kind": "review", "items": items})

    def _ai_chat_stream(self, payload):
        """stream:true 的 /api/ai/chat：上游 SSE 帧逐行透传（chunked 编码）。"""
        if not AI_BASE_URL or not AI_API_KEY:
            return self._send_json({"ok": False, "error": "服务端未配置 AI_BASE_URL / AI_API_KEY"})
        body = {
            "model": payload.get("model") or AI_DEFAULT_MODEL,
            "messages": payload.get("messages") or [],
            "temperature": payload.get("temperature", 0.7),
            "stream": True,
        }
        # 模型只允许 .env 白名单内的取值
        if AI_MODELS and body["model"] not in AI_MODELS:
            body["model"] = AI_DEFAULT_MODEL
        if payload.get("max_tokens"):
            body["max_tokens"] = int(payload["max_tokens"])
        req = urllib.request.Request(
            AI_BASE_URL + "/chat/completions",
            data=json.dumps(body).encode("utf-8"),
            headers={
                "Content-Type": "application/json",
                "Authorization": "Bearer " + AI_API_KEY,
            },
            method="POST",
        )
        try:
            resp = urllib.request.urlopen(req, timeout=AI_TIMEOUT)
        except urllib.error.HTTPError as e:
            detail = e.read().decode("utf-8", "ignore")[:500]
            return self._send_json({"ok": False, "error": "模型服务 HTTP %s: %s" % (e.code, detail)})
        except Exception as e:  # 超时/网络错误
            return self._send_json({"ok": False, "error": "模型调用失败: %s" % e})
        # 上游正常，开始 SSE 流
        self.send_response(200)
        self.send_header("Content-Type", "text/event-stream; charset=utf-8")
        self.send_header("Cache-Control", "no-cache")
        self.send_header("Transfer-Encoding", "chunked")
        self.end_headers()
        try:
            for raw_line in resp:
                line = raw_line.decode("utf-8", "ignore").rstrip("\r\n")
                if not line:
                    continue
                self._sse_chunk((line + "\n\n").encode("utf-8"))
        except Exception as e:
            # 流开始后才出错：以错误帧告知客户端再结束
            try:
                self._sse_chunk(("data: " + json.dumps(
                    {"error": "流式传输中断: %s" % e}, ensure_ascii=False) + "\n\n"
                ).encode("utf-8"))
            except Exception:
                pass
        finally:
            resp.close()
            try:
                self.wfile.write(b"0\r\n\r\n")  # chunked 结束标记
                self.wfile.flush()
            except Exception:
                pass

    def do_GET(self):
        path = urllib.parse.urlparse(self.path).path
        if path == "/api/constants":
            return self._send_json({
                "ok": True, "service": "sizheng-case-library",
                "aiConfigured": bool(AI_BASE_URL and AI_API_KEY),
                "models": AI_MODELS,
                "defaultModel": AI_DEFAULT_MODEL,
                "reviewEnabled": AI_REVIEW_ENABLED,
                "webSearch": bool(ANYSEARCH_API_KEY or TAVILY_API_KEY),
                "filesAuth": bool(APP_SECRET),
                "graph": GRAPH.configured,
            })
        if path == "/api/files":
            return self._send_json(api_files_list(auth_user(self)))
        if path == "/api/knowledge":
            return self._send_json(load_knowledge())
        if path == "/api/cases":
            qs = urllib.parse.parse_qs(urllib.parse.urlparse(self.path).query)
            res, status = api_cases_list(auth_user(self), qs)
            return self._send_json(res, status)
        if path == "/api/reviews":
            res, status = api_reviews(auth_user(self))
            return self._send_json(res, status)
        if path == "/api/admin/review-ledger":
            res, status = api_review_ledger(auth_user(self))
            return self._send_json(res, status)
        if path == "/api/favorites":
            res, status = api_favorites(auth_user(self))
            return self._send_json(res, status)
        if path == "/api/materials":
            qs = urllib.parse.parse_qs(urllib.parse.urlparse(self.path).query)
            res, status = api_materials_list(auth_user(self), qs)
            return self._send_json(res, status)
        if path == "/api/admin/watch/sources":
            res, status = api_watch_sources(auth_user(self))
            return self._send_json(res, status)
        if path == "/api/admin/watch/items":
            qs = urllib.parse.parse_qs(urllib.parse.urlparse(self.path).query)
            res, status = api_watch_items(auth_user(self), qs)
            return self._send_json(res, status)
        if path == "/api/contributions":
            res, status = api_contributions(auth_user(self))
            return self._send_json(res, status)
        if path == "/api/my/impact":
            res, status = api_my_impact(auth_user(self))
            return self._send_json(res, status)
        if path == "/api/my/prefs":
            res, status = api_my_prefs_get(auth_user(self))
            return self._send_json(res, status)
        if path == "/api/graph/overview":
            res, status = api_graph_overview(auth_user(self))
            return self._send_json(res, status)
        if path == "/api/graph/ego":
            qs = urllib.parse.parse_qs(urllib.parse.urlparse(self.path).query)
            res, status = api_graph_ego(auth_user(self), qs)
            return self._send_json(res, status)
        if path == "/api/graph/reverse":
            qs = urllib.parse.parse_qs(urllib.parse.urlparse(self.path).query)
            res, status = api_graph_reverse(auth_user(self), qs)
            return self._send_json(res, status)
        m = re.match(r"^/api/onlyoffice/config/([^/]+?)/?$", path)
        if m:
            res, status = api_oo_config(auth_user(self), urllib.parse.unquote(m.group(1)))
            return self._send_json(res, status)
        m = re.match(r"^/api/onlyoffice/file/([^/]+?)/?$", path)
        if m:
            data, err2, status = api_oo_file(self, urllib.parse.unquote(m.group(1)))
            if err2:
                return self._send_json(err2, status)
            self.send_response(200)
            self.send_header(
                "Content-Type",
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document")
            self.send_header("Cache-Control", "no-cache")
            self.send_header("Content-Length", str(len(data)))
            self.end_headers()
            self.wfile.write(data)
            return
        m = re.match(r"^/api/materials/([^/]+?)/?$", path)
        if m:
            res, status = api_material_get(auth_user(self), urllib.parse.unquote(m.group(1)))
            return self._send_json(res, status)
        m = re.match(r"^/api/cases/([^/]+)/attachments/([^/]+)/file/?$", path)
        if m:
            cid, aid = (urllib.parse.unquote(x) for x in m.groups())
            found, err, status = api_case_attachment_file(auth_user(self), cid, aid)
            if err:
                return self._send_json(err, status)
            data, ctype, name = found
            self.send_response(200)
            self.send_header("Content-Type", ctype)
            self.send_header("Content-Disposition", "inline; filename*=UTF-8''" + urllib.parse.quote(name))
            self.send_header("Content-Length", str(len(data)))
            self.end_headers()
            self.wfile.write(data)
            return
        m = re.match(r"^/api/cases/([^/]+?)(?:/(annotations|versions))?/?$", path)
        if m:
            cid = urllib.parse.unquote(m.group(1))
            res, status = api_case_get(auth_user(self), cid)
            if status == 200 and m.group(2):
                res = {"ok": True, m.group(2): res["case"][m.group(2)]}
            return self._send_json(res, status)
        if path.startswith("/api/files/"):
            rest = path[len("/api/files/"):]
            if rest.endswith("/text"):
                fid = urllib.parse.unquote(rest[:-len("/text")])
                res, status = api_file_text(auth_user(self), fid)
                return self._send_json(res, status)
            fid = urllib.parse.unquote(path.rsplit("/", 1)[1])
            found, err, status = api_file_download(auth_user(self), fid)
            if err:
                return self._send_json(err, status)
            data, ctype, name = found
            self.send_response(200)
            self.send_header("Content-Type", ctype)
            self.send_header("Content-Disposition",
                             "inline; filename*=UTF-8''" + urllib.parse.quote(name))
            self.send_header("Cache-Control", "no-cache")
            self.send_header("Content-Length", str(len(data)))
            self.end_headers()
            self.wfile.write(data)
            return
        return self._serve_static(path)

    def do_POST(self):
        path = urllib.parse.urlparse(self.path).path
        payload = self._read_json()
        if path == "/api/auth/login":
            return self._send_json(api_login(payload))
        if path == "/api/files":
            res, status = api_file_upload(auth_user(self), payload)
            return self._send_json(res, status)
        if path == "/api/knowledge/import":
            res, status = api_knowledge_import(auth_user(self), payload)
            return self._send_json(res, status)
        if path == "/api/ai/chat":
            if payload.get("stream"):
                return self._ai_chat_stream(payload)
            return self._send_json(ai_chat(payload))
        if path == "/api/search":
            res, status = api_search(auth_user(self), payload)
            return self._send_json(res, status)
        if path == "/api/cases":
            res, status = api_case_create(auth_user(self), payload)
            return self._send_json(res, status)
        if path == "/api/admin/reseed":
            res, status = api_reseed(auth_user(self))
            return self._send_json(res, status)
        if path == "/api/admin/materials/healthcheck":
            res, status = api_materials_healthcheck(auth_user(self), payload)
            return self._send_json(res, status)
        if path == "/api/materials":
            res, status = api_material_create(auth_user(self), payload)
            return self._send_json(res, status)
        m = re.match(r"^/api/cases/([^/]+)/attachments/?$", path)
        if m:
            res, status = api_case_attachment_add(
                auth_user(self), urllib.parse.unquote(m.group(1)), payload)
            return self._send_json(res, status)
        m = re.match(r"^/api/materials/([^/]+?)/favorite/?$", path)
        if m:
            res, status = api_mat_favorite(auth_user(self), urllib.parse.unquote(m.group(1)), True)
            return self._send_json(res, status)
        m = re.match(
            r"^/api/cases/([^/]+?)/(submit|withdraw|review|annotations|versions|favorite|like)/?$",
            path)
        if m:
            cid, sub = urllib.parse.unquote(m.group(1)), m.group(2)
            if sub in ("submit", "withdraw"):
                res, status = api_case_transition(auth_user(self), cid, sub, payload)
            elif sub == "review":
                res, status = api_case_transition(
                    auth_user(self), cid, (payload or {}).get("action") or "", payload)
            elif sub == "annotations":
                res, status = api_annotation_add(auth_user(self), cid, payload)
            elif sub == "versions":
                res, status = api_version_add(auth_user(self), cid, payload)
            elif sub == "favorite":
                res, status = api_favorite_set(auth_user(self), cid, True)
            else:
                res, status = api_like_set(auth_user(self), cid, True)
            return self._send_json(res, status)
        m = re.match(r"^/api/cases/([^/]+?)/versions/([^/]+?)/rollback/?$", path)
        if m:
            res, status = api_version_rollback(
                auth_user(self), urllib.parse.unquote(m.group(1)),
                urllib.parse.unquote(m.group(2)))
            return self._send_json(res, status)
        if path == "/api/ai/agent":
            return self._ai_agent(payload)
        if path == "/api/onlyoffice/callback":
            qs = urllib.parse.parse_qs(urllib.parse.urlparse(self.path).query)
            res, status = api_oo_callback(self, qs, payload)
            return self._send_json(res, status)
        if path == "/api/onlyoffice/insert":
            res, status = api_oo_insert(auth_user(self), payload)
            return self._send_json(res, status)
        if path == "/api/web-search":
            return self._send_json(web_search(payload))
        if path == "/api/graph/qa":
            res, status = api_graph_qa(auth_user(self), payload)
            return self._send_json(res, status)
        if path == "/api/admin/graph/rebuild":
            res, status = api_admin_graph_rebuild(auth_user(self))
            return self._send_json(res, status)
        if path == "/api/fetch-url":
            return self._send_json(fetch_url(payload))
        if path == "/api/admin/watch/sources":
            res, status = api_watch_source_create(auth_user(self), payload)
            return self._send_json(res, status)
        if path == "/api/admin/watch/run":
            res, status = api_watch_run(auth_user(self), payload)
            return self._send_json(res, status)
        m = re.match(r"^/api/admin/watch/items/([^/]+?)/import/?$", path)
        if m:
            res, status = api_watch_item_import(
                auth_user(self), urllib.parse.unquote(m.group(1)), payload)
            return self._send_json(res, status)
        if path == "/api/contributions":
            res, status = api_contribution_create(auth_user(self), payload)
            return self._send_json(res, status)
        m = re.match(r"^/api/contributions/([^/]+?)/review/?$", path)
        if m:
            res, status = api_contribution_review(
                auth_user(self), urllib.parse.unquote(m.group(1)), payload)
            return self._send_json(res, status)
        if path == "/api/export-docx":
            try:
                data, title = export_docx(payload)
            except Exception as e:
                return self._send_json({"ok": False, "error": "导出失败: %s" % e}, 500)
            fname = urllib.parse.quote(title + ".docx")
            self.send_response(200)
            self.send_header(
                "Content-Type",
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
            self.send_header("Content-Disposition",
                             "attachment; filename*=UTF-8''" + fname)
            self.send_header("Content-Length", str(len(data)))
            self.end_headers()
            self.wfile.write(data)
            return
        return self._send_json({"ok": False, "error": "not found"}, 404)

    def do_PUT(self):
        path = urllib.parse.urlparse(self.path).path
        if path == "/api/my/prefs":
            res, status = api_my_prefs_put(auth_user(self), self._read_json())
            return self._send_json(res, status)
        return self._send_json({"ok": False, "error": "not found"}, 404)

    def do_DELETE(self):
        path = urllib.parse.urlparse(self.path).path
        m = re.match(r"^/api/admin/watch/sources/([^/]+?)/?$", path)
        if m:
            res, status = api_watch_source_delete(
                auth_user(self), urllib.parse.unquote(m.group(1)))
            return self._send_json(res, status)
        if path.startswith("/api/files/"):
            fid = urllib.parse.unquote(path.rsplit("/", 1)[1])
            res, status = api_file_delete(auth_user(self), fid)
            return self._send_json(res, status)
        m = re.match(r"^/api/materials/([^/]+?)/favorite/?$", path)
        if m:
            res, status = api_mat_favorite(auth_user(self), urllib.parse.unquote(m.group(1)), False)
            return self._send_json(res, status)
        m = re.match(r"^/api/cases/([^/]+)/attachments/([^/]+)/?$", path)
        if m:
            cid, aid = (urllib.parse.unquote(x) for x in m.groups())
            res, status = api_case_attachment_delete(auth_user(self), cid, aid)
            return self._send_json(res, status)
        m = re.match(r"^/api/cases/([^/]+?)(?:/(favorite|like))?/?$", path)
        if m:
            cid, sub = urllib.parse.unquote(m.group(1)), m.group(2)
            if sub == "favorite":
                res, status = api_favorite_set(auth_user(self), cid, False)
            elif sub == "like":
                res, status = api_like_set(auth_user(self), cid, False)
            else:
                res, status = api_case_delete(auth_user(self), cid)
            return self._send_json(res, status)
        return self._send_json({"ok": False, "error": "not found"}, 404)

    def do_PATCH(self):
        path = urllib.parse.urlparse(self.path).path
        m = re.match(r"^/api/admin/watch/sources/([^/]+?)/?$", path)
        if m:
            res, status = api_watch_source_patch(
                auth_user(self), urllib.parse.unquote(m.group(1)), self._read_json())
            return self._send_json(res, status)
        m = re.match(r"^/api/admin/watch/items/([^/]+?)/?$", path)
        if m:
            res, status = api_watch_item_patch(
                auth_user(self), urllib.parse.unquote(m.group(1)), self._read_json())
            return self._send_json(res, status)
        if path.startswith("/api/files/"):
            fid = urllib.parse.unquote(path.rsplit("/", 1)[1])
            res, status = api_file_patch(auth_user(self), fid, self._read_json())
            return self._send_json(res, status)
        if path == "/api/materials":
            res, status = api_materials_batch(auth_user(self), self._read_json())
            return self._send_json(res, status)
        m = re.match(r"^/api/materials/([^/]+?)/?$", path)
        if m:
            res, status = api_material_patch(
                auth_user(self), urllib.parse.unquote(m.group(1)), self._read_json())
            return self._send_json(res, status)
        m = re.match(r"^/api/annotations/([^/]+?)/?$", path)
        if m:
            res, status = api_annotation_patch(
                auth_user(self), urllib.parse.unquote(m.group(1)), self._read_json())
            return self._send_json(res, status)
        m = re.match(r"^/api/cases/([^/]+?)/?$", path)
        if m:
            res, status = api_case_patch(
                auth_user(self), urllib.parse.unquote(m.group(1)), self._read_json())
            return self._send_json(res, status)
        return self._send_json({"ok": False, "error": "not found"}, 404)

    def _serve_static(self, path):
        if path in ("/", ""):
            path = "/index.html"
        safe = os.path.normpath(path).lstrip("/")
        if safe.startswith(".."):
            return self._send_json({"ok": False, "error": "forbidden"}, 403)
        full = os.path.join(STATIC_DIR, safe)
        if not os.path.isfile(full):
            full = os.path.join(STATIC_DIR, "index.html")  # SPA 兜底
        ctype = "text/html"
        if full.endswith(".js"):
            ctype = "application/javascript"
        elif full.endswith(".css"):
            ctype = "text/css"
        elif full.endswith(".svg"):
            ctype = "image/svg+xml"
        elif full.endswith(".png"):
            ctype = "image/png"
        elif full.endswith(".json"):
            ctype = "application/json"
        try:
            with open(full, "rb") as f:
                data = f.read()
        except OSError:
            return self._send_json({"ok": False, "error": "not found"}, 404)
        self.send_response(200)
        self.send_header("Content-Type", ctype + ("; charset=utf-8" if ctype.startswith(("text", "application/javascript")) else ""))
        self.send_header("Cache-Control", "no-cache")
        self.send_header("Content-Length", str(len(data)))
        self.end_headers()
        self.wfile.write(data)


def main():
    global CASEDB
    CASEDB = CaseDB(SQLITE_DB_PATH)
    seeded = CASEDB.seed(load_seed_cases(), load_seed_materials())
    CASEDB.seed_watch_sources(WATCH_SOURCES_PRESET)
    threading.Thread(target=watch_scheduler, daemon=True).start()
    backfilled = CASEDB.migrate_citation_evidence(_evidence_for_target)
    if backfilled:
        sys.stderr.write("[cite] 老引用证据回填 %d 条\n" % backfilled)
    if GRAPH.configured:
        db.GRAPH_HOOK = _graph_changed  # 案例/素材写路径增量同步图谱
        graph_rebuild()
    else:
        sys.stderr.write("[graph] 未配置 NEO4J_* 或缺少 neo4j 驱动"
                         "（图谱需 .venv/bin/python 运行），图谱功能关闭\n")
    if oo_configured():
        migrated = migrate_cases_docx()
        if migrated:
            sys.stderr.write("[oo] 案例 docx 迁移 %d 篇\n" % migrated)
    else:
        sys.stderr.write("[oo] 未配置 ONLYOFFICE_*，工作台使用兼容编辑器\n")
    port = int(sys.argv[1]) if len(sys.argv) > 1 else int(ENV.get("PROTOTYPE_PORT", "8080") or 8080)
    srv = ThreadingHTTPServer(("127.0.0.1", port), Handler)
    print("服务已启动: http://127.0.0.1:%d  (AI: %s, model=%s, 文件库: %s, 种子文件 %d 个, 业务库 %s%s)"
          % (port, "已配置" if (AI_BASE_URL and AI_API_KEY) else "未配置", AI_DEFAULT_MODEL,
             "鉴权开启" if APP_SECRET else "未配置 APP_SECRET", len(load_index()),
             os.path.relpath(SQLITE_DB_PATH, ROOT),
             "，首启灌入种子案例 %d 篇" % seeded if seeded else ""))
    try:
        srv.serve_forever()
    except KeyboardInterrupt:
        pass


if __name__ == "__main__":
    main()
